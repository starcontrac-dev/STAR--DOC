import os
from docx import Document

# Crear el directorio si no existe
base_dir = r"d:\GEMINI-DEV\copias\STAR-DOC-HTTPX-QWEN\copia\antigrabity\STAR-DOC-QWEN-1"
clausulas_dir = os.path.join(base_dir, "plantillas", "clausulas")
os.makedirs(clausulas_dir, exist_ok=True)

# 1. Crear la cláusula hija (SubDoc)
clausula_doc = Document()
clausula_doc.add_paragraph("CLÁUSULA SEXTA: RESOLUCIÓN DE CONTROVERSIAS. Cualquier disputa que surja en relación con el presente contrato, será sometida a consideración de un Tribunal de Arbitraje.")
clausula_path = os.path.join(clausulas_dir, "clausula_arbitraje.docx")
clausula_doc.save(clausula_path)
print(f"Creada clausula en: {clausula_path}")

# 2. Crear la plantilla principal (con SubDoc y {%p %})
main_doc = Document()
main_doc.add_heading("CONTRATO DE PRUEBA SUBDOC", level=1)
main_doc.add_paragraph("Este es un contrato para probar la inyección dinámica de SubDocs y el control de párrafos.")

# Prisma de control omitido para probar SubDoc directamente

# Prueba de inyección de cláusula (SubDoc)
main_doc.add_paragraph("A continuación la cláusula inyectada dinámicamente:")
main_doc.add_paragraph("{{ mi_clausula_dinamica }}")

main_path = os.path.join(base_dir, "plantillas", "subdoc_test.docx")
main_doc.save(main_path)
print(f"Creada plantilla principal en: {main_path}")
