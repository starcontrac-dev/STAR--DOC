"""
Analizador de Documentos NLP — FASE 7 STAR-DOC (Senior Refactor V2.1)

Herramienta global de análisis automático de documentos legales colombianos.
Refactorización orientada a la robustez en producción: Concurrencia segura,
manejo defensivo de excepciones, optimización de memoria, métricas de legibilidad
(Flesch-Szigriszt), clasificación automática de documentos y reranking semántico.

Archivo: analizador_documentos.py
"""

import re
import os
import logging
import asyncio
import threading
from typing import Dict, List, Any, Optional, Final, Set
from collections import Counter

from pydantic import BaseModel, Field, field_validator

# ============================================================
# Configuración y Constantes Globales
# ============================================================
logger = logging.getLogger(__name__)

MAX_TEXT_PROCESS_LENGTH: Final[int] = 150000  # Límite seguro para spaCy
MIN_TEXT_LENGTH_ANALYSIS: Final[int] = 20
MIN_TEXT_LENGTH_CLAUSES: Final[int] = 50

# ============================================================
# Carga Diferida (Lazy) Segura para Hilos (Thread-Safe Singletons)
# ============================================================

_nlp_es = None
_stopwords_es: Optional[Set[str]] = None
_nlp_lock = threading.Lock()
_stopwords_lock = threading.Lock()


def _get_nlp():
    """Carga lazy del modelo spaCy optimizado con seguridad para hilos."""
    global _nlp_es
    if _nlp_es is None:
        with _nlp_lock:
            if _nlp_es is None:  # Patrón Double-Checked Locking
                try:
                    import spacy
                    modelo = os.getenv("SPACY_MODEL_ES", "es_core_news_sm")
                    # Excluimos componentes innecesarios para optimizar RAM y CPU
                    # Mantenemos parser y attribute_ruler necesarios para noun_chunks
                    _nlp_es = spacy.load(modelo, exclude=["textcat"])
                    logger.info(f"✅ Modelo spaCy '{modelo}' cargado correctamente")
                except OSError as e:
                    logger.error(f"Error al cargar el modelo de spaCy: {e}")
                    raise RuntimeError(
                        f"Modelo de spaCy no encontrado. "
                        f"Ejecuta: python -m spacy download {os.getenv('SPACY_MODEL_ES', 'es_core_news_sm')}"
                    ) from e
    return _nlp_es


def _get_stopwords() -> Set[str]:
    """Carga lazy de stopwords en español de NLTK segura para hilos."""
    global _stopwords_es
    if _stopwords_es is None:
        with _stopwords_lock:
            if _stopwords_es is None:
                import nltk
                recursos = [
                    ("corpora/stopwords", "stopwords"),
                    ("tokenizers/punkt", "punkt"),
                    ("tokenizers/punkt_tab", "punkt_tab")
                ]
                for path, package in recursos:
                    try:
                        nltk.data.find(path)
                    except LookupError:
                        logger.info(f"Descargando recurso faltante de NLTK: {package}")
                        nltk.download(package, quiet=True)
                from nltk.corpus import stopwords
                _stopwords_es = set(stopwords.words("spanish"))
                logger.info("✅ Stopwords NLTK cargadas correctamente")
    return _stopwords_es


# ============================================================
# Schemas Pydantic para Validación de Entrada
# ============================================================

class AnalizarContratoInput(BaseModel):
    """Esquema de entrada para análisis de contratos."""
    texto: Optional[str] = Field(
        None,
        description="Texto directo del contrato a analizar (alternativa a file_path)"
    )
    file_path: Optional[str] = Field(
        None,
        description="Ruta al archivo del contrato (PDF, DOCX, TXT, MD)"
    )

    @field_validator("file_path")
    @classmethod
    def validar_extension(cls, v: Optional[str]) -> Optional[str]:
        if v is not None:
            extensiones_validas = (".pdf", ".docx", ".txt", ".md")
            if not v.lower().endswith(extensiones_validas):
                raise ValueError(
                    f"Formato no soportado. Extensiones válidas: {extensiones_validas}"
                )
        return v


class ExtraerEntidadesInput(BaseModel):
    """Esquema de entrada para extracción de entidades NER."""
    texto: str = Field(
        ...,
        min_length=MIN_TEXT_LENGTH_ANALYSIS,
        description="Texto del documento legal a analizar"
    )
    max_entidades: int = Field(
        50, ge=1, le=200, description="Número máximo de entidades a retornar por categoría"
    )


class DetectarClausulasInput(BaseModel):
    """Esquema de entrada para detección de cláusulas contractuales."""
    texto: str = Field(
        ...,
        min_length=MIN_TEXT_LENGTH_CLAUSES,
        description="Texto del contrato a analizar para detectar cláusulas"
    )


# ============================================================
# Extracción de Texto Robusta (Asíncrona)
# ============================================================

async def extraer_texto_pdf(file_path: str) -> str:
    """Extrae texto de un PDF usando manejo defensivo de errores."""
    def _extract() -> str:
        import pymupdf
        bloques_texto = []
        try:
            with pymupdf.open(file_path) as doc:
                for pagina in doc:
                    bloques = pagina.get_text("blocks")
                    for b in bloques:
                        # Indice 6 determina el tipo de bloque: 0 es texto, 1 es imagen
                        if len(b) > 6 and b[6] == 0:
                            texto_limpio = b[4].replace('\n', ' ').strip()
                            if texto_limpio:
                                bloques_texto.append(texto_limpio)
        except Exception as e:
            logger.error(f"Error procesando PDF {file_path}: {str(e)}")
            raise ValueError(f"No se pudo extraer texto del PDF: {str(e)}")
        
        return "\n\n".join(bloques_texto)

    return await asyncio.to_thread(_extract)


async def extraer_texto_docx(file_path: str) -> str:
    """Extrae texto de un archivo DOCX de forma asíncrona y segura."""
    def _extract() -> str:
        import docx
        partes = []
        try:
            doc = docx.Document(file_path)
            # Extraer párrafos
            partes.extend(p.text.strip() for p in doc.paragraphs if p.text.strip())
            
            # Extraer tablas optimizado
            for table in doc.tables:
                for row in table.rows:
                    fila_texto = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if fila_texto:
                        partes.append(fila_texto)
        except Exception as e:
            logger.error(f"Error procesando DOCX {file_path}: {str(e)}")
            raise ValueError(f"No se pudo extraer texto del DOCX: {str(e)}")
            
        return "\n".join(partes)

    return await asyncio.to_thread(_extract)


async def extraer_texto(file_path: str) -> str:
    """Delegador principal de extracción de archivos."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
    
    file_lower = file_path.lower()
    
    if file_lower.endswith(".pdf"):
        return await extraer_texto_pdf(file_path)
    elif file_lower.endswith(".docx"):
        return await extraer_texto_docx(file_path)
    elif file_lower.endswith((".txt", ".md")):
        def _read():
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        return await asyncio.to_thread(_read)
    else:
        raise ValueError(f"Formato de archivo no soportado: {file_path}")


# ============================================================
# Análisis NLP — Extracción y Lógica de Negocio
# ============================================================

# Expresiones regulares para leyes y referencias de Colombia expandidas y optimizadas
PATRONES_LEGALES = [
    # Captura Leyes, Decretos, Resoluciones, Circulares y Acuerdos de Colombia
    re.compile(r"\b(?:Ley|Decreto|Resolución|Circular|Acuerdo)\s+(?:Nro?\.?\s*)?\d+(?:\s*(?:de\s*|-)\s*\d{4}|\b)", re.I),
    # Captura Sentencias de Altas Cortes en Colombia (T-123 de 2023, C-007/18, SU-111/20, SL-456 de 2022, SC-789 de 2024, etc.)
    re.compile(r"\bSentencia\s+(?:[A-Z]{1,3}-\d+(?:\s*de\s*\d{4}|\s*/\s*\d{2,4}\b|\b))", re.I),
    # Códigos nacionales principales
    re.compile(r"\bCódigo\s+(?:Civil|de\s+Comercio|Sustantivo\s+del\s+Trabajo|General\s+del\s+Proceso|de\s+Procedimiento\s+Administrativo|Penal|de\s+Procedimiento\s+Penal)\b", re.I),
    # Constitución de Colombia
    re.compile(r"\b(?:Constitución\s+Política\s+(?:de\s+Colombia|de\s+1991|nacional)?|C\.P\.)\b", re.I),
    # Artículos específicos con ley/código referenciado (ej. Art. 61 del CST)
    re.compile(r"\b(?:Artículo|Art\.?)\s+\d+(?:\s*(?:literal\s*[a-z]|numeral\s*\d+)?(?:\s+(?:del|de\s+la)\s+(?:código|ley|decreto|cst|c\.c\.))?)?", re.I),
]


def extraer_entidades(texto: str, max_entidades: int = 50) -> Dict[str, List[str]]:
    """Extrae entidades nombradas usando spaCy y regex precompilados."""
    nlp = _get_nlp()
    
    # Prevenir OutOfMemory truncando textos excesivamente largos
    texto_seguro = texto[:MAX_TEXT_PROCESS_LENGTH]
    doc = nlp(texto_seguro)
    
    entidades: Dict[str, List[str]] = {
        "personas": [], "organizaciones": [], "fechas": [],
        "dinero": [], "ubicaciones": [], "legales": [],
    }
    
    vistos = {k: set() for k in entidades.keys()}
    
    # Mapeo de labels de spaCy a nuestras categorías
    label_map = {
        "PER": "personas", "ORG": "organizaciones", "DATE": "fechas",
        "TIME": "fechas", "MONEY": "dinero", "PERCENT": "dinero",
        "LOC": "ubicaciones", "GPE": "ubicaciones", "MISC": "legales"
    }
    
    for ent in doc.ents:
        texto_limpio = ent.text.strip()
        if not texto_limpio or ent.root.is_stop or ent.root.is_punct or len(texto_limpio) < 2:
            continue
            
        cat = label_map.get(ent.label_)
        if cat and texto_limpio not in vistos[cat]:
            entidades[cat].append(texto_limpio)
            vistos[cat].add(texto_limpio)
    
    # Enriquecer entidades legales con regex
    for patron in PATRONES_LEGALES:
        for match in patron.finditer(texto_seguro):
            match_limpio = match.group().strip()
            if match_limpio and match_limpio not in vistos["legales"]:
                entidades["legales"].append(match_limpio)
                vistos["legales"].add(match_limpio)
    
    # Limitar y retornar
    return {k: v[:max_entidades] for k, v in entidades.items()}


def segmentar_texto_legal(texto: str, chunk_size_words: int = 250) -> List[str]:
    """Divide el texto en fragmentos que respetan oraciones usando NLTK."""
    import nltk
    from nltk.tokenize import sent_tokenize
    
    # Forzamos la descarga asegurada por la función lazy
    _get_stopwords() 
    
    oraciones = sent_tokenize(texto, language="spanish")
    chunks, chunk_actual = [], []
    contador_palabras = 0
    
    for oracion in oraciones:
        palabras_oracion = len(oracion.split())
        if contador_palabras + palabras_oracion > chunk_size_words and chunk_actual:
            chunks.append(" ".join(chunk_actual))
            chunk_actual = [oracion]
            contador_palabras = palabras_oracion
        else:
            chunk_actual.append(oracion)
            contador_palabras += palabras_oracion
            
    if chunk_actual:
        chunks.append(" ".join(chunk_actual))
        
    return chunks


def extraer_temas_gensim(texto: str) -> List[str]:
    """Identifica tópicos latentes (LDA) manejando excepciones de vocabulario vacío."""
    try:
        from gensim import corpora
        from gensim.models import LdaModel
        from nltk.tokenize import word_tokenize
        
        stopwords = _get_stopwords()
        texto_limpio = re.sub(r'[^\w\s]', '', texto.lower())
        parrafos = [p for p in texto_limpio.split('\n') if len(p) > 50]
        
        if len(parrafos) < 3: 
            return []
            
        textos = [
            [word for word in word_tokenize(p) if word not in stopwords and word.isalpha()] 
            for p in parrafos
        ]
        
        # Prevenir fallos si el documento es puro texto irrelevante o números
        textos_filtrados = [t for t in textos if t]
        if not textos_filtrados:
            return []

        diccionario = corpora.Dictionary(textos_filtrados)
        if len(diccionario) == 0:
            return []

        corpus = [diccionario.doc2bow(text) for text in textos_filtrados]
        
        # Reducimos 'passes' para asegurar un tiempo de respuesta óptimo en backend
        lda = LdaModel(corpus, num_topics=2, id2word=diccionario, passes=5, random_state=42)
        
        return [" + ".join(word for word, _ in topic) for _, topic in lda.show_topics(formatted=False, num_words=3)]
        
    except Exception as e:
        logger.warning(f"Error no crítico en extracción de tópicos (Gensim): {e}")
        return []


def rerank_documentos_gensim(query: str, documentos: List[str], top_k: int = 3) -> List[Dict[str, Any]]:
    """
    Realiza un reranking (búsqueda por similitud semántica) entre un texto de consulta
    y una lista de documentos usando TF-IDF de Gensim, con un fallback
    seguro de similitud de Jaccard si Gensim falla o no está disponible.
    """
    if not documentos or not query:
        return []
        
    try:
        from gensim import corpora, models, similarities
        from nltk.tokenize import word_tokenize
        
        stopwords = _get_stopwords()
        
        def preprocesar(txt):
            limpio = re.sub(r'[^\w\s]', '', txt.lower())
            return [w for w in word_tokenize(limpio) if w not in stopwords and w.isalpha()]
            
        textos_corpus = [preprocesar(doc) for doc in documentos]
        texto_query = preprocesar(query)
        
        if not texto_query or not any(textos_corpus):
            raise ValueError("Vacío tras el preprocesamiento")
            
        diccionario = corpora.Dictionary(textos_corpus)
        if len(diccionario) == 0:
            raise ValueError("Diccionario vacío")
            
        corpus_bow = [diccionario.doc2bow(text) for text in textos_corpus]
        query_bow = diccionario.doc2bow(texto_query)
        
        # Modelo TF-IDF para ponderación
        tfidf = models.TfidfModel(corpus_bow)
        corpus_tfidf = tfidf[corpus_bow]
        query_tfidf = tfidf[query_bow]
        
        # Creación del índice de similitud
        indice = similarities.MatrixSimilarity(corpus_tfidf)
        similitudes = indice[query_tfidf]
        
        resultados = []
        for doc_idx, score in enumerate(similitudes):
            resultados.append({
                "indice": doc_idx,
                "score": round(float(score), 4),
                "texto_preview": documentos[doc_idx][:200] + "..." if len(documentos[doc_idx]) > 200 else documentos[doc_idx]
            })
            
        resultados.sort(key=lambda x: x["score"], reverse=True)
        return resultados[:top_k]
        
    except Exception as e:
        logger.warning(f"Gensim o preprocesamiento falló. Aplicando fallback de Jaccard. Detalle: {e}")
        # Fallback de similitud de Jaccard simple pero robusto
        def get_words(txt):
            return set(re.findall(r'\b\w{3,}\b', txt.lower()))
            
        query_words = get_words(query)
        resultados = []
        
        for idx, doc in enumerate(documentos):
            doc_words = get_words(doc)
            if not query_words or not doc_words:
                score = 0.0
            else:
                interseccion = query_words.intersection(doc_words)
                union = query_words.union(doc_words)
                score = len(interseccion) / len(union)
                
            resultados.append({
                "indice": idx,
                "score": round(score, 4),
                "texto_preview": doc[:200] + "..." if len(doc) > 200 else doc
            })
            
        resultados.sort(key=lambda x: x["score"], reverse=True)
        return resultados[:top_k]


# ============================================================
# Detección de Cláusulas y Riesgos Contractuales
# ============================================================

PATRONES_CLAUSULAS = {
    "objeto": re.compile(r"cláusula.*?objeto|objeto\s+(?:del\s+)?contrato|objeto\s+contractual", re.I),
    "salario": re.compile(r"salario|sueldo|remuneración|honorarios|contraprestación|pago\s+mensual|auxilio\s+de\s+transporte", re.I),
    "duracion": re.compile(r"duración|vigencia|plazo|término\s+(?:del\s+)?contrato|período\s+contractual|término\s+indefinido|período\s+de\s+prueba", re.I),
    "terminacion": re.compile(r"terminación|resolución|finalización|mutuo\s+desistimiento|causales?\s+de\s+terminación|despido", re.I),
    "confidencialidad": re.compile(r"confidencialidad|reserva|secreto\s+(?:profesional|comercial)|cláusula\s+de\s+(?:no\s+)?divulgación", re.I),
    "penalidades": re.compile(r"penalid|cláusula\s+penal|multa|sanción\s+contractual|pena\s+pecuniaria|perjuicios\s+estimados", re.I),
    "fuerza_mayor": re.compile(r"fuerza\s+mayor|caso\s+fortuito|circunstancias?\s+imprevisibles?", re.I),
    "propiedad_intelectual": re.compile(r"propiedad\s+intelectual|derechos?\s+de\s+autor|patente|marca|obra\s+intelectual|derechos\s+patrimoniales", re.I),
    "no_competencia": re.compile(r"no\s+competencia|exclusividad|pacto\s+de\s+no\s+competencia|restricción\s+competitiva", re.I),
    "jurisdiccion": re.compile(r"jurisdicción|competencia\s+judicial|tribunal|arbitraje|arbitral|juez\s+competente", re.I),
    "datos_personales": re.compile(r"datos?\s+personales?|ley\s+1581|protección\s+de\s+datos|habeas\s+data|autorización\s+de\s+tratamiento", re.I),
    "indemnizacion": re.compile(r"indemnizaci|resarcimiento|reparación\s+de\s+(?:daños|perjuicios)", re.I),
    "obligaciones": re.compile(r"obligacion(?:es)?\s+(?:del|de\s+las?\s+partes?)|deberes?\s+contractuales?", re.I),
    "garantias": re.compile(r"garantías?|póliza|seguro\s+de\s+cumplimiento|caución", re.I),
    "cesion": re.compile(r"cesión|ceder|transferir\s+(?:el\s+)?contrato|subcontrat", re.I),
    "ley_aplicable": re.compile(r"ley\s+aplicable|normativ(?:a|idad)\s+aplicable|legislación\s+colombiana|derecho\s+colombiano", re.I),
    "solucion_controversias": re.compile(r"solución\s+de\s+controversias?|mecanismo\s+de\s+resolución|conciliación|mediación|amigable\s+composición", re.I),
}

PATRONES_ABUSIVAS = {
    "exoneracion_responsabilidad": re.compile(r"exoner(?:a|ar|ación)\s+de\s+responsabilidad|no\s+se\s+hace\s+responsable|exime\s+de\s+toda\s+responsabilidad", re.I),
    "renuncia_derechos": re.compile(r"renuncia\s+(?:irrevocable\s+)?a\s+(?:reclamar|derechos|prestaciones|indemnización|auxilios)", re.I),
    "modificacion_unilateral": re.compile(r"modificar\s+unilateralmente|cambiar\s+sin\s+previo\s+aviso|derecho\s+de\s+modificar\s+de\s+forma\s+exclusiva", re.I),
    "arbitraje_obligatorio": re.compile(r"obligatoriamente\s+a\s+tribunal\s+de\s+arbitramento|renuncia\s+a\s+la\s+justicia\s+ordinaria", re.I),
    "mora_usuraria": re.compile(r"interés\s+del\s+\d+%\s+mensual|intereses?\s+moratorios?\s+del\s+\d+%", re.I)
}

# Constantes para comprobación lógica externa (Mantenidas para retrocompatibilidad)
CLAUSULAS_ESENCIALES = {
    "objeto", "duracion", "terminacion", "obligaciones", "jurisdiccion"
}

CLAUSULAS_RECOMENDADAS = {
    "confidencialidad", "penalidades", "datos_personales", 
    "fuerza_mayor", "ley_aplicable", "solucion_controversias"
}


def extraer_clausulas_contrato(texto: str) -> Dict[str, bool]:
    """Detecta la presencia de cláusulas verificando los regex precompilados."""
    clausulas = {nombre: bool(patron.search(texto)) for nombre, patron in PATRONES_CLAUSULAS.items()}
    clausulas.update({f"abusiva_{nombre}": bool(patron.search(texto)) for nombre, patron in PATRONES_ABUSIVAS.items()})
    return clausulas


def identificar_riesgos_contrato(clausulas: Dict[str, bool]) -> List[Dict[str, str]]:
    """Genera matrices de riesgos y fundamentos legales basados en las cláusulas."""
    riesgos = []
    
    # Reglas de negocio estructuradas con fundamentación de leyes colombianas actualizadas
    reglas_riesgos = [
        ("objeto", False, "alto", "objeto", "No se identifica el objeto del contrato de forma explícita.", "Código Civil Art. 1501"),
        ("terminacion", False, "alto", "terminación", "No se especifican las condiciones de terminación.", "C.S.T. Art. 61; C.Co. Art. 870"),
        ("obligaciones", False, "alto", "obligaciones", "No se detallan las obligaciones específicas de las partes.", "Código Civil Art. 1602"),
        ("abusiva_exoneracion_responsabilidad", True, "alto", "exoneración de responsabilidad", "Se detectó una posible cláusula abusiva de exoneración de responsabilidad.", "Ley 1480 de 2011 Art. 43"),
        ("abusiva_renuncia_derechos", True, "alto", "renuncia a derechos", "Se detectó una posible renuncia a derechos laborales mínimos irrenunciables o de consumo.", "C.S.T. Art. 14; Ley 1480 de 2011 Art. 43"),
        ("abusiva_modificacion_unilateral", True, "alto", "modificación unilateral", "Se detectó la facultad de modificación unilateral del contrato, prohibida en relaciones simétricas y de consumo.", "Código Civil Art. 1602; Ley 1480 de 2011"),
        ("abusiva_arbitraje_obligatorio", True, "medio", "arbitraje obligatorio", "Pacto arbitral que podría restringir el acceso del consumidor a la justicia ordinaria.", "Ley 1480 de 2011 Art. 43 Num. 8"),
        ("abusiva_mora_usuraria", True, "alto", "intereses de mora excesivos", "Posible estipulación de intereses moratorios que superan los límites legales de usura.", "Código de Comercio Art. 884; Código Civil Art. 1601"),
        ("confidencialidad", False, "medio", "confidencialidad", "Falta cláusula de confidencialidad y reserva de información.", "Ley 256 de 1996"),
        ("penalidades", False, "medio", "penalidades", "Sin cláusula penal establecida para caso de incumplimiento.", "Código Civil Art. 1592"),
        ("datos_personales", False, "medio", "datos personales", "No menciona política ni autorización de tratamiento de datos personales.", "Ley 1581 de 2012"),
        ("fuerza_mayor", False, "medio", "fuerza mayor", "No contempla exclusión de responsabilidad por fuerza mayor o caso fortuito.", "Código Civil Art. 64"),
        ("propiedad_intelectual", False, "bajo", "propiedad intelectual", "No se regula la propiedad intelectual ni los derechos de autor sobre las creaciones.", "Ley 23 de 1982"),
        ("cesion", False, "bajo", "cesión", "No regula la cesión del contrato a terceros.", "Código Civil Art. 1959"),
        ("garantias", False, "bajo", "garantías", "No se exigen garantías ni pólizas de cumplimiento.", "Código de Comercio Art. 1045"),
    ]

    for key, condicion, nivel, clausula, desc, base in reglas_riesgos:
        if clausulas.get(key) == condicion:
            riesgos.append({
                "nivel": nivel, "clausula": clausula, 
                "descripcion": desc, "fundamento_legal": base
            })

    # Regla compuesta
    if not clausulas.get("jurisdiccion") and not clausulas.get("solucion_controversias"):
        riesgos.append({
            "nivel": "medio", "clausula": "jurisdicción / solución",
            "descripcion": "No se define jurisdicción ni mecanismo de solución de controversias.",
            "fundamento_legal": "Ley 1563 de 2012"
        })
        
    return riesgos


def calcular_score_salud(clausulas: Dict[str, bool], riesgos: List[Dict[str, str]]) -> Dict[str, Any]:
    """Calcula la salud general del contrato usando ponderaciones de riesgo."""
    if not clausulas:
        return {"score": 0, "estado": "Crítico"}
        
    penalidades = {"alto": 20, "medio": 10, "bajo": 5}
    puntaje_base = 100 - sum(penalidades.get(r["nivel"], 0) for r in riesgos)
    score_final = max(0, min(100, puntaje_base))
    
    estado = "Excelente" if score_final >= 80 else \
             "Aceptable" if score_final >= 60 else \
             "Riesgoso" if score_final >= 40 else "Crítico"
        
    return {"score": score_final, "estado": estado}


# ============================================================
# Análisis y Estadísticas del Texto (Avanzadas)
# ============================================================

def contar_silabas_espanol(palabra: str) -> int:
    """
    Estima el número de sílabas de una palabra en español de forma lingüísticamente precisa.
    Maneja diptongos, triptongos, hiatos y la 'h' intermedia muda.
    """
    palabra = palabra.lower().strip()
    # Eliminar puntuación y caracteres no alfabéticos
    palabra = re.sub(r'[^\wáéóíúü]', '', palabra)
    if not palabra:
        return 0
        
    # Casos especiales de 'y' final precedida de vocal (actúa como vocal 'i')
    if palabra.endswith('y') and len(palabra) > 1 and palabra[-2] in 'aeiouáéóíúü':
        palabra = palabra[:-1] + 'i'
        
    # Eliminar 'h' intermedia entre vocales para no romper diptongos/hiatos
    palabra = re.sub(r'([aeiouáéóíúü])h([aeiouáéóíúü])', r'\1\2', palabra)
    
    # Expresión regular para identificar núcleos vocálicos en español:
    # 1. Triptongos: cerrada átona + abierta + cerrada átona (ej: iai, uai)
    # 2. Diptongos: abierta + cerrada átona, cerrada átona + abierta, o dos cerradas distintas
    # 3. Monoptongos: cualquier vocal individual
    # Las cerradas con tilde (í, ú) actúan como tónicas/fuertes (formando hiato)
    nucleos = re.findall(
        r'([iuü][aeoáéó][iuü]|[aeoáéó][iuü]|[iuü][aeoáéó]|[iuü][iuü]|[aeiouáéóíúü])', 
        palabra
    )
    
    return max(1, len(nucleos))


def calcular_flesch_szigriszt(total_palabras: int, total_silabas: int, total_oraciones: int) -> float:
    """Calcula el Índice de Flesch-Szigriszt (INFLESZ) para legibilidad en español."""
    if total_palabras == 0 or total_oraciones == 0:
        return 0.0
    prom_silabas_palabra = total_silabas / total_palabras
    prom_palabras_oracion = total_palabras / total_oraciones
    score = 206.84 - 60.0 * prom_silabas_palabra - 1.02 * prom_palabras_oracion
    return round(score, 2)


def clasificar_legibilidad_szigriszt(score: float) -> str:
    """Clasifica el nivel de legibilidad del Índice Flesch-Szigriszt en español."""
    if score >= 80:
        return "Muy Fácil"
    elif score >= 70:
        return "Algo Fácil"
    elif score >= 60:
        return "Normal (Estilo adecuado)"
    elif score >= 50:
        return "Algo Difícil"
    elif score >= 30:
        return "Difícil (Estilo técnico / legal complejo)"
    else:
        return "Muy Difícil (Estilo científico / legal de muy alta complejidad)"


def clasificar_documento_legal(texto: str, conceptos_clave: List[str]) -> str:
    """
    Clasifica automáticamente el tipo de documento legal utilizando coincidencia
    de patrones, entidades y términos del derecho colombiano.
    """
    texto_lower = texto.lower()
    
    # 1. Acción de Tutela
    tutela_keywords = ["acción de tutela", "accion de tutela", "juez de tutela", "derechos fundamentales", "vulneración", "pretensiones", "tutelar"]
    if any(kw in texto_lower for kw in tutela_keywords) and "hechos" in texto_lower:
        return "Acción de Tutela"
        
    # 2. Derecho de Petición
    peticion_keywords = ["derecho de petición", "derecho de peticion", "artículo 23 de la constitución", "respetuosamente solicito", "peticionario", "solicitud respetuosa"]
    if any(kw in texto_lower for kw in peticion_keywords) and "solicito" in texto_lower:
        return "Derecho de Petición"
        
    # 3. Contrato de Trabajo
    if "contrato individual de trabajo" in texto_lower or "código sustantivo del trabajo" in texto_lower or (
        "contrato de trabajo" in texto_lower and any(w in texto_lower for w in ["empleador", "trabajador", "salario", "cst"])
    ):
        return "Contrato de Trabajo"
        
    # 4. Contrato de Arrendamiento
    arrendamiento_keywords = ["arrendador", "arrendatario", "canon de arrendamiento", "inmueble", "arrendar", "contrato de arrendamiento", "servicios públicos", "destinación"]
    if any(kw in texto_lower for kw in arrendamiento_keywords) and ("arrendador" in texto_lower and "arrendatario" in texto_lower):
        return "Contrato de Arrendamiento"
        
    # 5. Contrato de Prestación de Servicios
    servicios_keywords = ["contrato de prestación de servicios", "prestacion de servicios", "contratante", "contratista", "honorarios", "objeto del contrato", "independencia del contratista"]
    if any(kw in texto_lower for kw in servicios_keywords) and ("contratante" in texto_lower and "contratista" in texto_lower):
        return "Contrato de Prestación de Servicios"
        
    # 6. Contrato de Compraventa
    compraventa_keywords = ["compraventa", "vendedor", "comprador", "precio de venta", "cosa vendida", "tradición", "saneamiento por evicción"]
    if any(kw in texto_lower for kw in compraventa_keywords) and ("vendedor" in texto_lower and "comprador" in texto_lower):
        return "Contrato de Compraventa"
        
    # 7. Poder Especial
    poder_keywords = ["otorgo poder", "confiero poder", "poder especial", "apoderado", "mandante", "mandatario", "representación judicial"]
    if any(kw in texto_lower for kw in poder_keywords) and ("poder" in texto_lower and "representar" in texto_lower):
        return "Poder Especial"
        
    # 8. Sentencia Judicial
    sentencia_keywords = ["administrando justicia", "en nombre de la república", "fallo", "resuelve", "corte constitucional", "corte suprema de justicia", "consejo de estado", "magistrado ponente"]
    if any(kw in texto_lower for kw in sentencia_keywords) and ("resuelve" in texto_lower or "decide" in texto_lower):
        return "Sentencia Judicial"
        
    # 9. Fallback usando sustantivos y conceptos clave
    if "contrato" in texto_lower or "cláusula" in texto_lower or "partes" in texto_lower:
        return "Otro Contrato / Acuerdo"
        
    return "Otro Documento Legal"


def calcular_estadisticas(texto: str) -> Dict[str, Any]:
    """Calcula métricas estructurales e indicadores de legibilidad del documento (Flesch-Szigriszt)."""
    nlp = _get_nlp()
    doc = nlp(texto[:15000]) # Usamos una muestra representativa
    
    palabras = texto.split()
    total_palabras = len(palabras)
    oraciones = list(doc.sents)
    total_oraciones = len(oraciones)
    parrafos = [p for p in texto.split("\n") if p.strip()]
    
    # Métricas de Calidad de Texto
    palabras_unicas = len(set(palabras))
    densidad_lexica = round((palabras_unicas / total_palabras * 100), 2) if total_palabras > 0 else 0
    prom_palabras_oracion = round(total_palabras / total_oraciones, 1) if total_oraciones > 0 else 0
    
    # Estimar total de sílabas en el texto (muestra)
    # Para optimizar rendimiento en CPU, calculamos sílabas en una muestra de palabras
    muestra_palabras = palabras[:2000]
    total_silabas_muestra = sum(contar_silabas_espanol(p) for p in muestra_palabras)
    
    # Escalar el conteo de sílabas al total de palabras
    if len(muestra_palabras) > 0:
        total_silabas_estimado = int((total_silabas_muestra / len(muestra_palabras)) * total_palabras)
    else:
        total_silabas_estimado = 0
        
    score_szigriszt = calcular_flesch_szigriszt(total_palabras, total_silabas_estimado, total_oraciones)
    legibilidad_szigriszt = clasificar_legibilidad_szigriszt(score_szigriszt)

    conceptos_clave = [
        chunk.text for chunk in doc.noun_chunks 
        if not chunk.root.is_stop and len(chunk.text) > 4
    ]
    top_conceptos = [c[0] for c in Counter(conceptos_clave).most_common(10)]
    
    # Clasificación automática del tipo de documento legal
    tipo_documento = clasificar_documento_legal(texto, top_conceptos)
    
    return {
        "total_palabras": total_palabras,
        "total_caracteres": len(texto),
        "total_oraciones": total_oraciones,
        "total_parrafos": len(parrafos),
        "promedio_palabras_por_oracion": prom_palabras_oracion,
        "densidad_lexica_pct": densidad_lexica,
        "legibilidad_flesch_szigriszt": score_szigriszt,
        "nivel_legibilidad": legibilidad_szigriszt,
        "tipo_documento_clasificado": tipo_documento,
        "conceptos_juridicos_principales": top_conceptos[:5]
    }


# ============================================================
# Core API / Endpoints de Servicio
# ============================================================

async def analizar_contrato(
    texto: Optional[str] = None,
    file_path: Optional[str] = None
) -> Dict[str, Any]:
    """Orquestador principal que consolida todo el análisis del documento."""
    if texto is None and file_path:
        texto = await extraer_texto(file_path)
    elif texto is None:
        raise ValueError("Se requiere proporcionar 'texto' o 'file_path'.")
    
    if len(texto.strip()) < MIN_TEXT_LENGTH_ANALYSIS:
        return {
            "error": "El texto es demasiado corto para realizar un análisis significativo.",
            "longitud": len(texto.strip()),
            "minimo_requerido": MIN_TEXT_LENGTH_ANALYSIS
        }
    
    logger.info(f"📄 Analizando documento de {len(texto)} caracteres...")
    
    # Ejecución paralela de tareas CPU-bound
    entidades_task = asyncio.to_thread(extraer_entidades, texto)
    stats_task = asyncio.to_thread(calcular_estadisticas, texto)
    topicos_task = asyncio.to_thread(extraer_temas_gensim, texto[:30000])
    chunks_task = asyncio.to_thread(segmentar_texto_legal, texto[:50000])
    
    entidades, stats, topicos, chunks_inteligentes = await asyncio.gather(
        entidades_task, stats_task, topicos_task, chunks_task
    )
    
    clausulas = extraer_clausulas_contrato(texto)
    riesgos = identificar_riesgos_contrato(clausulas)
    
    # Mapa de Normatividad protegido
    normas_detectadas = entidades.get("legales", [])
    mapa_normatividad = []
    if normas_detectadas:
        try:
            # Intento de integración segura (Falla silenciosamente si el módulo no existe)
            from app.core.tools.buscador_jurisprudencia import validar_normatividad_documento
            mapa_normatividad = await validar_normatividad_documento(normas_detectadas[:5]) 
        except Exception as e:
            logger.debug(f"Integración de jurisprudencia omitida/no disponible: {e}")
            
    if topicos:
        stats["topicos_latentes_gensim"] = topicos
        
    # Construcción de la respuesta compatible hacia atrás
    resumen_riesgos = {
        "altos": sum(1 for r in riesgos if r["nivel"] == "alto"),
        "medios": sum(1 for r in riesgos if r["nivel"] == "medio"),
        "bajos": sum(1 for r in riesgos if r["nivel"] == "bajo"),
    }
    
    stats.update({
        "total_entidades": sum(len(v) for v in entidades.values()),
        "clausulas_detectadas": sum(1 for v in clausulas.values() if v),
        "clausulas_faltantes": sum(1 for v in clausulas.values() if not v),
        "total_clausulas_evaluadas": len(clausulas),
        "riesgos_identificados": len(riesgos),
    })
    
    recomendaciones = []
    if resumen_riesgos["altos"] > 0:
        recomendaciones.append(f"⚠️ CRÍTICO: {resumen_riesgos['altos']} riesgos altos detectados. Revise los elementos esenciales.")
    
    recomendaciones.extend([f"🔴 {r['descripcion']}" for r in riesgos if r["nivel"] == "alto"])
    recomendaciones.extend([f"🟡 {r['descripcion']}" for r in riesgos if r["nivel"] == "medio"])
    
    if not riesgos:
        recomendaciones.append("✅ El contrato cumple con las cláusulas principales base.")
    recomendaciones.append("📋 IMPORTANTE: Análisis generado mediante NLP. Requiere validación de un profesional del derecho.")
    
    clausulas_totales = len(clausulas)
    clausulas_activas = sum(1 for v in clausulas.values() if v)
    score_completitud = round((clausulas_activas / clausulas_totales * 100), 1) if clausulas_totales else 100
    salud = calcular_score_salud(clausulas, riesgos)
    
    resultado = {
        "entidades": entidades,
        "clausulas": {k: v for k, v in clausulas.items() if not k.startswith("abusiva_")},
        "clausulas_abusivas": {k.replace("abusiva_", ""): v for k, v in clausulas.items() if k.startswith("abusiva_")},
        "mapa_normatividad": mapa_normatividad,
        "riesgos": riesgos,
        "resumen_riesgos": resumen_riesgos,
        "estadisticas": stats,
        "score_completitud": score_completitud,
        "salud_contractual": salud,
        "chunks_recomendados": len(chunks_inteligentes),
        "recomendaciones": recomendaciones,
    }
    
    logger.info(f"✅ Análisis completado. Score Salud: {salud['score']}%")
    return resultado


def extraer_entidades_documento(texto: str, max_entidades: int = 50) -> Dict[str, Any]:
    """Wrapper sincrónico compatible."""
    if len(texto.strip()) < MIN_TEXT_LENGTH_ANALYSIS:
        return {"error": "Texto demasiado corto para análisis de entidades."}
    
    entidades = extraer_entidades(texto, max_entidades)
    total = sum(len(v) for v in entidades.values())
    
    return {
        "entidades": entidades,
        "total_entidades": total,
        "categorias": {k: len(v) for k, v in entidades.items()},
        "message": f"Se detectaron {total} entidades en {sum(1 for v in entidades.values() if v)} categorías."
    }


def detectar_clausulas_documento(texto: str) -> Dict[str, Any]:
    """Wrapper sincrónico compatible."""
    if len(texto.strip()) < MIN_TEXT_LENGTH_CLAUSES:
        return {"error": "Texto demasiado corto para análisis de cláusulas."}
    
    clausulas = extraer_clausulas_contrato(texto)
    riesgos = identificar_riesgos_contrato(clausulas)
    
    detectadas = [k for k, v in clausulas.items() if v]
    faltantes = [k for k, v in clausulas.items() if not v]
    
    return {
        "clausulas": clausulas,
        "clausulas_detectadas": detectadas,
        "clausulas_faltantes": faltantes,
        "riesgos": riesgos,
        "total_detectadas": len(detectadas),
        "total_faltantes": len(faltantes),
        "score_completitud": round(len(detectadas) / len(clausulas) * 100, 1) if clausulas else 100,
        "message": f"Se detectaron {len(detectadas)} cláusulas y {len(riesgos)} riesgos."
    }