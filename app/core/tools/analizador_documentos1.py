"""
Analizador de Documentos NLP — FASE 7 STAR-DOC (Senior Refactor)

Herramienta global de análisis automático de documentos legales colombianos
usando spaCy 3.8 + NLTK 3.9 + Gensim 4.3 para extracción de entidades, 
detección de cláusulas contractuales, identificación de riesgos, modelado 
de tópicos y generación de informes estructurados.

Decisión Arquitectónica:
  - NO se usa lexnlp (incompatible con Python 3.13)
  - Se usa spaCy (es_core_news_sm/md) para NER + análisis morfológico
  - Se usa Gensim para modelado de tópicos latentes
  - Asyncio para I/O no bloqueante en lectura de archivos
  - Regex precompilado y optimizado para léxico jurídico colombiano

Autor: Equipo STAR-DOC
Fecha de Refactor: 2026-04-20
"""

import re
import os
import logging
import asyncio
from typing import Dict, List, Any, Optional
from pydantic import BaseModel, Field, field_validator

logger = logging.getLogger(__name__)

# ============================================================
# Carga diferida (lazy) y Cacheo de modelos NLP
# ============================================================

_nlp_es = None
_stopwords_es = None


def _get_nlp():
    """Carga lazy del modelo spaCy optimizado."""
    global _nlp_es
    if _nlp_es is None:
        try:
            import spacy
            # Permite inyectar un modelo más pesado en prod (ej. es_core_news_lg)
            modelo = os.getenv("SPACY_MODEL_ES", "es_core_news_sm")
            # Desactivamos textcat si no se usa para ahorrar RAM y ganar velocidad
            _nlp_es = spacy.load(modelo, exclude=["textcat"])
            logger.info(f"✅ Modelo spaCy '{modelo}' cargado correctamente")
        except OSError:
            raise RuntimeError(
                f"Modelo de spaCy no encontrado. "
                f"Ejecuta: python -m spacy download {os.getenv('SPACY_MODEL_ES', 'es_core_news_sm')}"
            )
    return _nlp_es


def _get_stopwords():
    """Carga lazy de stopwords en español de NLTK."""
    global _stopwords_es
    if _stopwords_es is None:
        import nltk
        try:
            nltk.data.find("corpora/stopwords")
            nltk.data.find("tokenizers/punkt")
        except LookupError:
            nltk.download("stopwords", quiet=True)
            nltk.download("punkt", quiet=True)
            nltk.download("punkt_tab", quiet=True)
        from nltk.corpus import stopwords
        _stopwords_es = set(stopwords.words("spanish"))
        logger.info("✅ Stopwords NLTK cargadas correctamente")
    return _stopwords_es


# ============================================================
# Schemas Pydantic para validación de entrada
# ============================================================

class AnalizarContratoInput(BaseModel):
    """Esquema de entrada para análisis de contratos."""
    texto: Optional[str] = Field(
        None,
        description="Texto directo del contrato a analizar (alternativa a file_path)"
    )
    file_path: Optional[str] = Field(
        None,
        description="Ruta al archivo del contrato (PDF, DOCX, TXT, MD)"
    )

    @field_validator("file_path")
    @classmethod
    def validar_extension(cls, v: Optional[str]) -> Optional[str]:
        """Valida que el archivo tenga extensión soportada."""
        if v is not None:
            extensiones_validas = (".pdf", ".docx", ".txt", ".md")
            if not v.lower().endswith(extensiones_validas):
                raise ValueError(
                    f"Formato no soportado. Extensiones válidas: {extensiones_validas}"
                )
        return v


class ExtraerEntidadesInput(BaseModel):
    """Esquema de entrada para extracción de entidades NER."""
    texto: str = Field(
        ...,
        min_length=20,
        description="Texto del documento legal a analizar (mínimo 20 caracteres)"
    )
    max_entidades: int = Field(
        50,
        ge=1,
        le=200,
        description="Número máximo de entidades a retornar por categoría"
    )


class DetectarClausulasInput(BaseModel):
    """Esquema de entrada para detección de cláusulas contractuales."""
    texto: str = Field(
        ...,
        min_length=50,
        description="Texto del contrato a analizar para detectar cláusulas (mín 50 chars)"
    )


# ============================================================
# Extracción de texto desde archivos (Asíncrono)
# ============================================================

async def extraer_texto_pdf(file_path: str) -> str:
    """
    Extrae texto de un archivo PDF usando PyMuPDF (fitz).
    Utiliza el método 'blocks' para no romper oraciones entre líneas.
    """
    def _extract() -> str:
        import pymupdf
        doc = pymupdf.open(file_path)
        bloques_texto = []
        for pagina in doc:
            bloques = pagina.get_text("blocks")
            for b in bloques:
                if b[6] == 0:  # Validamos que sea bloque de texto, no imagen
                    texto_limpio = b[4].replace('\n', ' ').strip()
                    if texto_limpio:
                        bloques_texto.append(texto_limpio)
        doc.close()
        return "\n\n".join(bloques_texto)

    return await asyncio.to_thread(_extract)


async def extraer_texto_docx(file_path: str) -> str:
    """Extrae texto de un archivo DOCX de forma asíncrona."""
    def _extract() -> str:
        import docx
        doc = docx.Document(file_path)
        partes = []
        
        # Extraer párrafos
        for p in doc.paragraphs:
            if p.text.strip():
                partes.append(p.text.strip())
        
        # Extraer tablas
        for table in doc.tables:
            for row in table.rows:
                fila_texto = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if fila_texto:
                    partes.append(fila_texto)
        
        return "\n".join(partes)

    return await asyncio.to_thread(_extract)


async def extraer_texto(file_path: str) -> str:
    """Extrae texto de un archivo delegando al formato correcto."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
    
    file_lower = file_path.lower()
    
    if file_lower.endswith(".pdf"):
        return await extraer_texto_pdf(file_path)
    elif file_lower.endswith(".docx"):
        return await extraer_texto_docx(file_path)
    elif file_lower.endswith((".txt", ".md")):
        def _read():
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        return await asyncio.to_thread(_read)
    else:
        raise ValueError(f"Formato no soportado: {file_path}")


# ============================================================
# Análisis NLP — Extracción de Entidades y Tópicos
# ============================================================

def extraer_entidades(texto: str, max_entidades: int = 50) -> Dict[str, List[str]]:
    """Extrae entidades nombradas usando spaCy y regex precompilados."""
    nlp = _get_nlp()
    
    # Limitar texto para evitar desbordamiento de memoria
    texto_limitado = texto[:100000]
    doc = nlp(texto_limitado)
    
    entidades = {
        "personas": [],
        "organizaciones": [],
        "fechas": [],
        "dinero": [],
        "ubicaciones": [],
        "legales": [],
    }
    
    # Set temporal para búsquedas O(1)
    vistos = {k: set() for k in entidades.keys()}
    
    for ent in doc.ents:
        texto_limpio = ent.text.strip()
        
        # Filtro nativo spaCy
        if not texto_limpio or ent.root.is_stop or ent.root.is_punct or len(texto_limpio) < 2:
            continue
            
        label = ent.label_
        cat = None
        
        if label == "PER": cat = "personas"
        elif label == "ORG": cat = "organizaciones"
        elif label in ("DATE", "TIME"): cat = "fechas"
        elif label in ("MONEY", "PERCENT"): cat = "dinero"
        elif label in ("LOC", "GPE"): cat = "ubicaciones"
        elif label == "MISC": cat = "legales"
        
        if cat and texto_limpio not in vistos[cat]:
            entidades[cat].append(texto_limpio)
            vistos[cat].add(texto_limpio)
    
    # Enriquecer entidades legales con regex precompilados
    patrones_legales = [
        re.compile(r"(?:Ley|Decreto|Resolución|Acuerdo|Circular)\s+\d+\s*(?:de\s+\d{4})?", re.I),
        re.compile(r"(?:Artículo|Art\.?)\s+\d+(?:\s+(?:del|de la)\s+[\w\s]+)?", re.I),
        re.compile(r"Sentencia\s+[A-Z]-\d+(?:\s+de\s+\d{4})?", re.I),
        re.compile(r"Código\s+(?:Civil|Penal|de Comercio|Sustantivo del Trabajo|General del Proceso)", re.I),
        re.compile(r"Constitución\s+(?:Política\s+)?(?:de\s+(?:Colombia|1991))?", re.I),
    ]
    
    for patron in patrones_legales:
        for match in patron.finditer(texto_limitado):
            match_limpio = match.group().strip()
            if match_limpio and match_limpio not in vistos["legales"]:
                entidades["legales"].append(match_limpio)
                vistos["legales"].add(match_limpio)
    
    # Limitar resultados por categoría
    for key in entidades:
        entidades[key] = entidades[key][:max_entidades]
    
    return entidades


def segmentar_texto_legal(texto: str, chunk_size_words: int = 250) -> List[str]:
    """
    Segmentación Inteligente (Chunking) usando NLTK.
    Divide el texto en fragmentos que respetan los límites de las oraciones (sentencias legales),
    ideal para inyectar contexto semánticamente coherente en la Bóveda RAG o LLM.
    """
    import nltk
    from nltk.tokenize import sent_tokenize
    
    try:
        nltk.data.find("tokenizers/punkt")
        nltk.data.find("tokenizers/punkt_tab")
    except LookupError:
        nltk.download("punkt", quiet=True)
        nltk.download("punkt_tab", quiet=True)
        
    oraciones = sent_tokenize(texto, language="spanish")
    
    chunks = []
    chunk_actual = []
    contador_palabras = 0
    
    for oracion in oraciones:
        palabras_oracion = len(oracion.split())
        
        if contador_palabras + palabras_oracion > chunk_size_words and chunk_actual:
            chunks.append(" ".join(chunk_actual))
            chunk_actual = [oracion]
            contador_palabras = palabras_oracion
        else:
            chunk_actual.append(oracion)
            contador_palabras += palabras_oracion
            
    if chunk_actual:
        chunks.append(" ".join(chunk_actual))
        
    return chunks


def extraer_temas_gensim(texto: str) -> List[str]:
    """Identifica los tópicos latentes del documento usando LDA (Gensim)."""
    try:
        from gensim import corpora
        from gensim.models import LdaModel
        import nltk
        from nltk.tokenize import word_tokenize
        
        stopwords = _get_stopwords()
        texto_limpio = re.sub(r'[^\w\s]', '', texto.lower())
        parrafos = [p for p in texto_limpio.split('\n') if len(p) > 50]
        
        if len(parrafos) < 3: 
            return []
            
        textos = [
            [word for word in word_tokenize(p) if word not in stopwords and word.isalpha()] 
            for p in parrafos
        ]
        
        diccionario = corpora.Dictionary(textos)
        corpus = [diccionario.doc2bow(text) for text in textos]
        lda = LdaModel(corpus, num_topics=2, id2word=diccionario, passes=10, random_state=42)
        
        temas = []
        for idx, topic in lda.show_topics(formatted=False, num_words=3):
            palabras = [word for word, prob in topic]
            temas.append(" + ".join(palabras))
            
        return temas
    except ImportError:
        logger.warning("Gensim no está instalado. Omitiendo modelado de tópicos.")
        return []
    except Exception as e:
        logger.error(f"Error en Gensim: {e}")
        return []


def rerank_documentos_gensim(query: str, documentos: List[str], top_k: int = 3) -> List[Dict[str, Any]]:
    """
    Realiza un reranking (búsqueda por similitud semántica) entre un texto de consulta
    y una lista de documentos usando TF-IDF o LSI de Gensim.
    Ideal para encontrar 'Documentos Relacionados' en la Bóveda RAG.
    """
    if not documentos or not query:
        return []
        
    try:
        from gensim import corpora, models, similarities
        from nltk.tokenize import word_tokenize
        
        stopwords = _get_stopwords()
        
        def preprocesar(txt):
            limpio = re.sub(r'[^\w\s]', '', txt.lower())
            return [w for w in word_tokenize(limpio) if w not in stopwords and w.isalpha()]
            
        textos_corpus = [preprocesar(doc) for doc in documentos]
        texto_query = preprocesar(query)
        
        if not texto_query or not any(textos_corpus):
            return []
            
        diccionario = corpora.Dictionary(textos_corpus)
        corpus_bow = [diccionario.doc2bow(text) for text in textos_corpus]
        query_bow = diccionario.doc2bow(texto_query)
        
        # Usar modelo TF-IDF para ponderar las palabras
        tfidf = models.TfidfModel(corpus_bow)
        corpus_tfidf = tfidf[corpus_bow]
        query_tfidf = tfidf[query_bow]
        
        # Opcional: LSI para semántica latente si hay muchos documentos
        # lsi = models.LsiModel(corpus_tfidf, id2word=diccionario, num_topics=min(10, len(documentos)))
        
        indice = similarities.MatrixSimilarity(corpus_tfidf)
        similitudes = indice[query_tfidf]
        
        # Ordenar por similitud
        resultados = []
        for doc_idx, score in enumerate(similitudes):
            resultados.append({
                "indice": doc_idx,
                "score": float(score),
                "texto_preview": documentos[doc_idx][:200] + "..."
            })
            
        resultados.sort(key=lambda x: x["score"], reverse=True)
        return resultados[:top_k]
        
    except ImportError:
        logger.warning("Gensim no está instalado. Omitiendo Reranking.")
        return []
    except Exception as e:
        logger.error(f"Error en Reranking Gensim: {e}")
        return []

# ============================================================
# Detección de Cláusulas y Riesgos Contractuales
# ============================================================

# Precompilado para mayor rendimiento
PATRONES_CLAUSULAS = {
    "objeto": re.compile(r"cláusula.*?objeto|objeto\s+(?:del\s+)?contrato|objeto\s+contractual", re.I),
    "salario": re.compile(r"salario|sueldo|remuneración|honorarios|contraprestación|pago\s+mensual", re.I),
    "duracion": re.compile(r"duración|vigencia|plazo|término\s+(?:del\s+)?contrato|período\s+contractual", re.I),
    "terminacion": re.compile(r"terminación|resolución|finalización|mutuo\s+desistimiento|causales?\s+de\s+terminación", re.I),
    "confidencialidad": re.compile(r"confidencialidad|reserva|secreto\s+(?:profesional|comercial)|cláusula\s+de\s+(?:no\s+)?divulgación", re.I),
    "penalidades": re.compile(r"penalid|cláusula\s+penal|multa|sanción\s+contractual|pena\s+pecuniaria", re.I),
    "fuerza_mayor": re.compile(r"fuerza\s+mayor|caso\s+fortuito|circunstancias?\s+imprevisibles?", re.I),
    "propiedad_intelectual": re.compile(r"propiedad\s+intelectual|derechos?\s+de\s+autor|patente|marca|obra\s+intelectual", re.I),
    "no_competencia": re.compile(r"no\s+competencia|exclusividad|pacto\s+de\s+no\s+competencia|restricción\s+competitiva", re.I),
    "jurisdiccion": re.compile(r"jurisdicción|competencia\s+judicial|tribunal|arbitraje|arbitral|juez\s+competente", re.I),
    "datos_personales": re.compile(r"datos?\s+personales?|ley\s+1581|protección\s+de\s+datos|habeas\s+data|autorización\s+de\s+tratamiento", re.I),
    "indemnizacion": re.compile(r"indemnizaci|resarcimiento|reparación\s+de\s+(?:daños|perjuicios)", re.I),
    "obligaciones": re.compile(r"obligacion(?:es)?\s+(?:del|de\s+las?\s+partes?)|deberes?\s+contractuales?", re.I),
    "garantias": re.compile(r"garantías?|póliza|seguro\s+de\s+cumplimiento|caución", re.I),
    "cesion": re.compile(r"cesión|ceder|transferir\s+(?:el\s+)?contrato|subcontrat", re.I),
    "ley_aplicable": re.compile(r"ley\s+aplicable|normativ(?:a|idad)\s+aplicable|legislación\s+colombiana|derecho\s+colombiano", re.I),
    "solucion_controversias": re.compile(r"solución\s+de\s+controversias?|mecanismo\s+de\s+resolución|conciliación|mediación|amigable\s+composición", re.I),
}

PATRONES_ABUSIVAS = {
    "exoneracion_responsabilidad": re.compile(r"exoner(?:a|ar|ación)\s+de\s+responsabilidad|no\s+se\s+hace\s+responsable", re.I),
    "renuncia_derechos": re.compile(r"renuncia\s+(?:irrevocable\s+)?a\s+(?:reclamar|derechos|prestaciones)", re.I),
    "modificacion_unilateral": re.compile(r"modificar\s+unilateralmente|cambiar\s+sin\s+previo\s+aviso", re.I)
}

# Sets restaurados para validaciones lógicas externas
CLAUSULAS_ESENCIALES = {
    "objeto", "duracion", "terminacion", "obligaciones", "jurisdiccion"
}

CLAUSULAS_RECOMENDADAS = {
    "confidencialidad", "penalidades", "datos_personales", 
    "fuerza_mayor", "ley_aplicable", "solucion_controversias"
}


def extraer_clausulas_contrato(texto: str) -> Dict[str, bool]:
    """Detecta la presencia de cláusulas clave usando regex precompilado."""
    clausulas = {}
    for nombre, patron in PATRONES_CLAUSULAS.items():
        clausulas[nombre] = bool(patron.search(texto))
        
    for nombre, patron in PATRONES_ABUSIVAS.items():
        clausulas[f"abusiva_{nombre}"] = bool(patron.search(texto))
        
    return clausulas


def identificar_riesgos_contrato(clausulas: Dict[str, bool]) -> List[Dict[str, str]]:
    """Identifica riesgos legales basados en cláusulas faltantes o problemáticas."""
    riesgos = []
    
    # === RIESGOS ALTOS (cláusulas esenciales) ===
    
    if not clausulas.get("objeto"):
        riesgos.append({
            "nivel": "alto",
            "clausula": "objeto",
            "descripcion": "No se identifica el objeto del contrato. Es un elemento esencial sin el cual el contrato puede ser inexistente.",
            "fundamento_legal": "Código Civil Colombiano, Art. 1501 — Elementos esenciales del contrato"
        })
    
    if not clausulas.get("terminacion"):
        riesgos.append({
            "nivel": "alto",
            "clausula": "terminación",
            "descripcion": "No se especifican condiciones de terminación. Genera indefinición contractual y riesgo de litigio.",
            "fundamento_legal": "Código Sustantivo del Trabajo Art. 61; Código de Comercio Art. 870"
        })
    
    if not clausulas.get("obligaciones"):
        riesgos.append({
            "nivel": "alto",
            "clausula": "obligaciones",
            "descripcion": "No se detallan las obligaciones de las partes. Dificulta exigir el cumplimiento contractual.",
            "fundamento_legal": "Código Civil Art. 1602 — El contrato es ley para las partes"
        })
        
    # === RIESGOS EXTREMOS (Cláusulas abusivas detectadas) ===
    if clausulas.get("abusiva_exoneracion_responsabilidad"):
        riesgos.append({
            "nivel": "alto",
            "clausula": "exoneración de responsabilidad",
            "descripcion": "Se detectó cláusula de exoneración de responsabilidad. Podría ser ineficaz de pleno derecho.",
            "fundamento_legal": "Estatuto del Consumidor (Ley 1480 de 2011) Art. 43"
        })
        
    if clausulas.get("abusiva_renuncia_derechos"):
        riesgos.append({
            "nivel": "alto",
            "clausula": "renuncia a derechos",
            "descripcion": "Se detectó cláusula de renuncia a derechos. Si es laboral o de consumo, es nula.",
            "fundamento_legal": "C.S.T. Art. 14 — Irrenunciabilidad; Ley 1480 Art. 43"
        })
        
    if clausulas.get("abusiva_modificacion_unilateral"):
        riesgos.append({
            "nivel": "alto",
            "clausula": "modificación unilateral",
            "descripcion": "El contrato permite modificación unilateral. Puede considerarse desequilibrio injustificado.",
            "fundamento_legal": "Código Civil; Ley 1480 de 2011"
        })
    
    # === RIESGOS MEDIOS (cláusulas recomendadas) ===
    
    if not clausulas.get("confidencialidad"):
        riesgos.append({
            "nivel": "medio",
            "clausula": "confidencialidad",
            "descripcion": "Falta cláusula de confidencialidad. Puede exponer información comercial sensible.",
            "fundamento_legal": "Ley 256 de 1996 (Competencia desleal); Código de Comercio Art. 19 num. 4"
        })
    
    if not clausulas.get("penalidades"):
        riesgos.append({
            "nivel": "medio",
            "clausula": "penalidades",
            "descripcion": "Sin cláusula penal. Dificulta cuantificar y exigir reparación por incumplimiento.",
            "fundamento_legal": "Código Civil Art. 1592 — Cláusula penal"
        })
    
    if not clausulas.get("jurisdiccion") and not clausulas.get("solucion_controversias"):
        riesgos.append({
            "nivel": "medio",
            "clausula": "jurisdicción / solución de controversias",
            "descripcion": "No se define jurisdicción ni mecanismo de solución de controversias. Puede generar conflictos de competencia.",
            "fundamento_legal": "Ley 1563 de 2012 (Arbitraje); Código General del Proceso Art. 28"
        })
    
    if not clausulas.get("datos_personales"):
        riesgos.append({
            "nivel": "medio",
            "clausula": "datos personales",
            "descripcion": "No menciona protección de datos personales. Puede constituir infracción regulatoria.",
            "fundamento_legal": "Ley 1581 de 2012 — Protección de Datos Personales; Decreto 1377 de 2013"
        })
    
    if not clausulas.get("fuerza_mayor"):
        riesgos.append({
            "nivel": "medio",
            "clausula": "fuerza mayor",
            "descripcion": "No contempla fuerza mayor o caso fortuito. Las partes quedan expuestas ante eventos imprevisibles.",
            "fundamento_legal": "Código Civil Art. 1 (fuerza mayor); Código de Comercio Art. 992"
        })
    
    # === RIESGOS BAJOS (buenas prácticas) ===
    
    if not clausulas.get("propiedad_intelectual"):
        riesgos.append({
            "nivel": "bajo",
            "clausula": "propiedad intelectual",
            "descripcion": "No se regula la propiedad intelectual. Recomendable en contratos con creación de obras o software.",
            "fundamento_legal": "Ley 23 de 1982 — Derechos de Autor; Decisión Andina 486"
        })
    
    if not clausulas.get("cesion"):
        riesgos.append({
            "nivel": "bajo",
            "clausula": "cesión",
            "descripcion": "No se regula la posibilidad de cesión del contrato. Podría ser cedido sin consentimiento.",
            "fundamento_legal": "Código Civil Art. 1959 — Cesión de derechos"
        })
    
    if not clausulas.get("garantias"):
        riesgos.append({
            "nivel": "bajo",
            "clausula": "garantías",
            "descripcion": "No se exigen garantías de cumplimiento. Considerar póliza o caución.",
            "fundamento_legal": "Código de Comercio Art. 1045 y ss. — Garantías"
        })
    
    return riesgos


def calcular_score_salud(clausulas: Dict[str, bool], riesgos: List[Dict[str, str]]) -> Dict[str, Any]:
    """
    Calcula un 'Score de Salud Contractual' (0-100%).
    Penaliza fuertemente riesgos altos (faltan cláusulas esenciales).
    """
    if not clausulas:
        return {"score": 0, "estado": "Crítico"}
        
    puntaje_base = 100
    
    for riesgo in riesgos:
        if riesgo["nivel"] == "alto":
            puntaje_base -= 20  # Penalización severa
        elif riesgo["nivel"] == "medio":
            puntaje_base -= 10
        elif riesgo["nivel"] == "bajo":
            puntaje_base -= 5
            
    score_final = max(0, min(100, puntaje_base))
    
    if score_final >= 80:
        estado = "Excelente"
    elif score_final >= 60:
        estado = "Aceptable"
    elif score_final >= 40:
        estado = "Riesgoso"
    else:
        estado = "Crítico"
        
    return {"score": score_final, "estado": estado}


# ============================================================
# Estadísticas textuales del documento
# ============================================================

def calcular_estadisticas(texto: str) -> Dict[str, Any]:
    """Calcula métricas y extrae conceptos jurídicos clave."""
    nlp = _get_nlp()
    doc = nlp(texto[:15000])
    
    palabras = texto.split()
    oraciones = list(doc.sents)
    parrafos = [p for p in texto.split("\n") if p.strip()]
    
    # Noun Chunks: Extraemos entidades compuestas para dar contexto semántico
    conceptos_clave = [
        chunk.text for chunk in doc.noun_chunks 
        if not chunk.root.is_stop and len(chunk.text) > 4
    ]
    from collections import Counter
    top_conceptos = [c[0] for c in Counter(conceptos_clave).most_common(5)]
    
    return {
        "total_palabras": len(palabras),
        "total_caracteres": len(texto),
        "total_oraciones": len(oraciones),
        "total_parrafos": len(parrafos),
        "promedio_palabras_por_oracion": round(len(palabras) / len(oraciones) if oraciones else 0, 1),
        "conceptos_juridicos_principales": top_conceptos
    }


# ============================================================
# Función principal de análisis — Punto de entrada para el agente
# ============================================================

async def analizar_contrato(
    texto: Optional[str] = None,
    file_path: Optional[str] = None
) -> Dict[str, Any]:
    """Orquestador principal asíncrono para el análisis completo."""
    
    if texto is None and file_path:
        texto = await extraer_texto(file_path)
    elif texto is None:
        raise ValueError("Se requiere proporcionar 'texto' o 'file_path' para el análisis.")
    
    if len(texto.strip()) < 20:
        return {
            "error": "El texto es demasiado corto para realizar un análisis significativo.",
            "longitud": len(texto.strip()),
            "minimo_requerido": 20
        }
    
    logger.info(f"📄 Analizando documento de {len(texto)} caracteres...")
    
    # Ejecutamos procesos intensivos de CPU en hilos separados
    entidades = await asyncio.to_thread(extraer_entidades, texto)
    clausulas = extraer_clausulas_contrato(texto)
    riesgos = identificar_riesgos_contrato(clausulas)
    stats = await asyncio.to_thread(calcular_estadisticas, texto)
    
    # Mapa de Normatividad
    normas_detectadas = entidades.get("legales", [])
    mapa_normatividad = []
    if normas_detectadas:
        try:
            from app.core.tools.buscador_jurisprudencia import validar_normatividad_documento
            mapa_normatividad = await validar_normatividad_documento(normas_detectadas[:5]) # Limitar a 5 para no saturar
        except ImportError:
            logger.warning("No se pudo importar validar_normatividad_documento.")
        except Exception as e:
            logger.error(f"Error al validar normatividad: {e}")
            
    # Modelado de tópicos opcional
    topicos = await asyncio.to_thread(extraer_temas_gensim, texto[:30000])
    if topicos:
        stats["topicos_latentes_gensim"] = topicos
        
    stats.update({
        "total_entidades": sum(len(v) for v in entidades.values()),
        "clausulas_detectadas": sum(1 for v in clausulas.values() if v),
        "clausulas_faltantes": sum(1 for v in clausulas.values() if not v),
        "total_clausulas_evaluadas": len(clausulas),
        "riesgos_identificados": len(riesgos),
    })
    
    resumen_riesgos = {
        "altos": len([r for r in riesgos if r["nivel"] == "alto"]),
        "medios": len([r for r in riesgos if r["nivel"] == "medio"]),
        "bajos": len([r for r in riesgos if r["nivel"] == "bajo"]),
    }
    
    recomendaciones = []
    
    if resumen_riesgos["altos"] > 0:
        recomendaciones.append(
            f"⚠️ CRÍTICO: Se detectaron {resumen_riesgos['altos']} riesgos altos. "
            "Revisar elementos esenciales del contrato ANTES de firmar."
        )
    
    for riesgo in riesgos:
        if riesgo["nivel"] == "alto":
            recomendaciones.append(f"🔴 {riesgo['descripcion']}")
        elif riesgo["nivel"] == "medio":
            recomendaciones.append(f"🟡 {riesgo['descripcion']}")
    
    if not riesgos:
        recomendaciones.append(
            "✅ El contrato parece contener las cláusulas principales. "
            "Se recomienda validación final por abogado especializado."
        )
    
    recomendaciones.append(
        "📋 IMPORTANTE: Este análisis es automatizado con NLP. "
        "Se recomienda SIEMPRE la revisión por un profesional del derecho."
    )
    
    clausulas_encontradas = sum(1 for v in clausulas.values() if v)
    score_completitud = round(
        (clausulas_encontradas / len(clausulas) * 100) if clausulas else 100, 1
    )
    
    salud = calcular_score_salud(clausulas, riesgos)
    
    # Añadimos chunking inteligente al resultado
    chunks_inteligentes = await asyncio.to_thread(segmentar_texto_legal, texto[:50000])
    
    # Reranking (ejemplo si tuviéramos base de documentos en bóveda)
    # rerank_docs = rerank_documentos_gensim(texto, documentos_boveda)
    
    resultado = {
        "entidades": entidades,
        "clausulas": {k: v for k, v in clausulas.items() if not k.startswith("abusiva_")},
        "clausulas_abusivas": {k.replace("abusiva_", ""): v for k, v in clausulas.items() if k.startswith("abusiva_")},
        "mapa_normatividad": mapa_normatividad,
        "riesgos": riesgos,
        "resumen_riesgos": resumen_riesgos,
        "estadisticas": stats,
        "score_completitud": score_completitud,
        "salud_contractual": salud,
        "chunks_recomendados": len(chunks_inteligentes),
        "recomendaciones": recomendaciones,
    }
    
    logger.info(f"✅ Análisis completado. Score Salud: {salud['score']}%")
    
    return resultado


def extraer_entidades_documento(texto: str, max_entidades: int = 50) -> Dict[str, Any]:
    """Wrapper sincrónico para extracción de entidades."""
    if len(texto.strip()) < 20:
        return {"error": "Texto demasiado corto para análisis de entidades."}
    
    entidades = extraer_entidades(texto, max_entidades)
    total = sum(len(v) for v in entidades.values())
    
    return {
        "entidades": entidades,
        "total_entidades": total,
        "categorias": {k: len(v) for k, v in entidades.items()},
        "message": f"Se detectaron {total} entidades en {len([k for k, v in entidades.items() if v])} categorías."
    }


def detectar_clausulas_documento(texto: str) -> Dict[str, Any]:
    """Wrapper sincrónico para detección de cláusulas."""
    if len(texto.strip()) < 50:
        return {"error": "Texto demasiado corto para análisis de cláusulas."}
    
    clausulas = extraer_clausulas_contrato(texto)
    riesgos = identificar_riesgos_contrato(clausulas)
    
    detectadas = [k for k, v in clausulas.items() if v]
    faltantes = [k for k, v in clausulas.items() if not v]
    
    return {
        "clausulas": clausulas,
        "clausulas_detectadas": detectadas,
        "clausulas_faltantes": faltantes,
        "riesgos": riesgos,
        "total_detectadas": len(detectadas),
        "total_faltantes": len(faltantes),
        "score_completitud": round(len(detectadas) / len(clausulas) * 100, 1) if clausulas else 100,
        "message": f"Se detectaron {len(detectadas)} cláusulas. {len(riesgos)} riesgos."
    }