from app.models.ipfs_audit import IPFSAudit
import os
import io
import logging
from typing import Optional
import aiofiles
import qrcode
import fitz  # PyMuPDF
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.document_ipfs import DocumentIPFS
from app.services.ipfs_service import IPFSService
from app.services.crypto_engine import CryptoEngine, DocClassification
from app.core.config import settings

logger = logging.getLogger(__name__)

class IPFSIntegrationService:
    """
    Servicio de integración de alto nivel para anclar documentos generados
    a IPFS y estamparlos con un QR Code de verificación.
    """

    @staticmethod
    async def anchor_and_stamp(
        file_path: str,
        classification: DocClassification = DocClassification.PUBLIC,
        user_id: Optional[int] = None,
        session: Optional[AsyncSession] = None,
        document_id: Optional[int] = None,
        signature_request_id: Optional[int] = None
    ) -> DocumentIPFS:
        """
        1. Lee el archivo limpio y calcula su Hash SHA-256 original.
        2. Sube una copia preliminar local a Kubo para obtener el CID original limpio.
        3. Estampa la evidencia digital con la certificación correspondiente:
           - Si es PDF: Código QR e información legible al pie de página.
           - Si es DOCX: Sección de certificación formal al final del documento.
        4. Sube la versión estampada definitiva a IPFS (Kubo + Pinata híbrido si corresponde).
           Si la subida falla, encola la tarea en la base de datos de reintentos para no interrumpir el flujo.
        5. Registra el CID definitivo (o provisional) en la base de datos (PostgreSQL).
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"El archivo {file_path} no existe.")

        filename = os.path.basename(file_path)

        # 1. Leer el archivo limpio y calcular su SHA-256 original
        import hashlib
        async with aiofiles.open(file_path, "rb") as f:
            file_data = await f.read()
        
        sha256_original = hashlib.sha256(file_data).hexdigest()

        # Asegurar tipo de clasificación a string plano limpio para evitar fallos de Enum/SQLAlchemy en runtime
        class_str = classification.value if hasattr(classification, "value") else str(classification)
        class_str = class_str.strip().lower()

        # Generar llave de encriptación compartida si no es público
        encryption_key = None
        if class_str != "public":
            from cryptography.hazmat.primitives.ciphers.aead import AESGCM
            encryption_key = AESGCM.generate_key(bit_length=256)

        # 2. Subida rápida local a Kubo para obtener el CID original (inmutable)
        logger.info(f"Calculando CID original para evidencia: {filename}")
        try:
            if class_str != "public":
                # Encriptamos la versión limpia antes de subirla
                envelope_clean = CryptoEngine.encrypt_with_envelope(file_data, key=encryption_key)
                clean_upload_data = envelope_clean["encrypted_data"]
                clean_filename = f"{filename}.enc"
            else:
                clean_upload_data = file_data
                clean_filename = f"clean_{filename}"

            kubo_pre = await IPFSService.upload_to_kubo(clean_filename, clean_upload_data)
            cid_original = kubo_pre["cid"]
            logger.info(f"CID original obtenido de Kubo: {cid_original}")
        except Exception as e:
            logger.error(f"Error al obtener el CID original de Kubo: {e}")
            cid_original = f"sha256-{sha256_original}"

        # 3. Estampar metadatos en el archivo local
        if filename.lower().endswith(".pdf"):
            try:
                await IPFSIntegrationService._stamp_pdf_with_qr(file_path, cid_original, sha256_original)
                logger.info(f"PDF estampado exitosamente con el CID original: {cid_original}")
            except Exception as e:
                logger.error(f"Fallo al estampar el PDF con QR: {e}")
        elif filename.lower().endswith(".docx"):
            try:
                await IPFSIntegrationService._stamp_docx_with_text(file_path, cid_original, sha256_original)
                logger.info(f"DOCX estampado exitosamente con el CID original: {cid_original}")
            except Exception as e:
                logger.error(f"Fallo al estampar el DOCX con texto: {e}")

        # 4. Leer la versión estampada final
        async with aiofiles.open(file_path, "rb") as f:
            stamped_file_data = await f.read()
        
        file_size_stamped = os.path.getsize(file_path)

        # 5. Subir la versión ESTAMPADA a IPFS
        logger.info(f"Subiendo versión estampada final de {filename} a IPFS (Clasificación: {class_str})")
        
        ipfs_cid = None
        pinned_kubo = False
        pinned_pinata = False
        gateway_url = None
        is_encrypted = class_str != "public"
        
        try:
            result = await IPFSService.secure_upload(
                file_name=filename,
                file_data=stamped_file_data,
                classification=classification,
                encryption_key=encryption_key,
                metadata={
                    "source": "star-doc-generator",
                    "user_id": str(user_id),
                    "sha256_original": sha256_original,
                    "cid_original": cid_original
                }
            )
            ipfs_cid = result["cid"]
            pinned_kubo = result.get("kubo") is not None
            pinned_pinata = result.get("pinata") is not None
            gateway_url = result.get("gateway_url")
            is_encrypted = result["is_encrypted"]
            logger.info(f"Subida exitosa de evidencia a IPFS CID: {ipfs_cid}")
        except Exception as ipfs_err:
            import uuid
            import shutil
            logger.warning(f"⚠️ Error al subir a IPFS en caliente: {ipfs_err}. Encolando en segundo plano...")
            
            # Crear ID temporal único para no romper unicidad
            ipfs_cid = f"pending-ipfs-{uuid.uuid4().hex[:12]}"
            
            # Copiar el archivo a un directorio persistente dedicado a reintentos
            # Esto evita que se borre si proviene de directorios temporales (como en batch)
            pending_dir = os.path.join(settings.OUTPUT_DIR, "ipfs_pending")
            os.makedirs(pending_dir, exist_ok=True)
            pending_filename = f"{uuid.uuid4().hex}_{filename}"
            persist_path = os.path.join(pending_dir, pending_filename)
            try:
                shutil.copy2(file_path, persist_path)
                logger.info(f"Archivo copiado a ruta de reintentos persistente: {persist_path}")
            except Exception as copy_err:
                logger.error(f"No se pudo copiar el archivo a la ruta de reintentos: {copy_err}")
                persist_path = file_path  # Fallback al original
            
            # Crear tarea en la base de datos de reintentos
            from app.models.ipfs_pending_task import IPFSPendingTask
            pending_task = IPFSPendingTask(
                file_path=persist_path,
                classification=classification.value,
                user_id=user_id,
                document_id=document_id,
                signature_request_id=signature_request_id,
                retry_count=0,
                status="pending",
                last_error=str(ipfs_err)
            )
            
            if session is None:
                from app.database import async_session_maker
                async with async_session_maker() as new_session:
                    new_session.add(pending_task)
                    await new_session.commit()
            else:
                session.add(pending_task)
                # No hacemos commit aquí si el session es externo y será commiteado después
            
        # 6. Registrar en Base de Datos (PostgreSQL)
        import mimetypes
        mime = mimetypes.guess_type(filename)[0] or "application/octet-stream"

        doc_record = DocumentIPFS(
            document_id=document_id,
            user_id=user_id,
            ipfs_cid=ipfs_cid,
            ipfs_cid_original=cid_original,
            sha256_original=sha256_original,
            classification=classification.value,
            is_encrypted=is_encrypted,
            original_filename=filename,
            file_size_bytes=file_size_stamped,
            mime_type=mime,
            pinned_kubo=pinned_kubo,
            pinned_pinata=pinned_pinata,
            gateway_url=gateway_url,
        )

        if encryption_key and doc_record.user_id:
            doc_record.encryption_key_encrypted = CryptoEngine.encrypt_document_key(
                encryption_key, doc_record.user_id
            )

        if session is None:
            from app.database import async_session_maker
            async with async_session_maker() as new_session:
                new_session.add(doc_record)
                await new_session.commit()
                await new_session.refresh(doc_record)
        else:
            session.add(doc_record)
            await session.commit()
            await session.refresh(doc_record)
        
        logger.info(f"Evidencia registrada en BD con ID: {doc_record.id} y CID: {ipfs_cid}")
        return doc_record

    @staticmethod
    async def _stamp_pdf_with_qr(pdf_path: str, cid_original: str, sha256_original: str):
        """
        Inyecta un código QR y metadatos de auditoría al pie de página de la última hoja del PDF.
        Si la última página no es de firmas o auditoría, se crea una nueva página dedicada al final
        para evitar pisar la información original del contrato.
        """
        def _sync_stamp():
            doc = fitz.open(pdf_path)
            last_page = doc[-1]
            page_text = last_page.get_text()
            
            # Verificar si ya es una página especial de firmas o certificado
            is_special_page = (
                "ANEXO DE FIRMAS" in page_text or 
                "CERTIFICADO DE FIRMA" in page_text or 
                "CADENA DE CUSTODIA" in page_text
            )
            
            page_width = last_page.rect.width
            page_height = last_page.rect.height
            
            # Generar QR
            verify_url = f"https://ipfs.io/ipfs/{cid_original}"
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_M,
                box_size=10,
                border=2,
            )
            qr.add_data(verify_url)
            qr.make(fit=True)
            img = qr.make_image(fill_color="black", back_color="white")
            
            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr)
            img_byte_arr.seek(0)
            
            if is_special_page:
                # Estampar en la esquina inferior de la página especial existente
                qr_size = 75
                margin = 25
                
                rect = fitz.Rect(
                    page_width - qr_size - margin, 
                    page_height - qr_size - margin, 
                    page_width - margin, 
                    page_height - margin
                )
                last_page.insert_image(rect, stream=img_byte_arr.read())
                
                # Inyectar metadatos legibles en la esquina inferior izquierda
                cert_box_rect = fitz.Rect(
                    margin + 5, 
                    page_height - margin - 35, 
                    page_width - qr_size - margin - 15, 
                    page_height - margin
                )
                
                text_content = (
                    f"CERTIFICACION DE EVIDENCIA STAR-DOC (IPFS)\n"
                    f"SHA-256 Original: {sha256_original}\n"
                    f"Dirección IPFS: https://ipfs.io/ipfs/{cid_original}"
                )
                last_page.insert_textbox(
                    cert_box_rect,
                    text_content,
                    fontsize=6.0,
                    fontname="Helvetica",
                    color=(0.25, 0.28, 0.35)
                )
            else:
                # Crear una nueva página dedicada al depósito de evidencia IPFS
                cert_page = doc.new_page()
                page_width = cert_page.rect.width
                page_height = cert_page.rect.height
                
                # Dibujar banner superior corporativo
                cert_page.draw_rect(fitz.Rect(20, 20, page_width - 20, 75), color=None, fill=(0.01, 0.24, 0.43), overlay=True)
                cert_page.insert_textbox(
                    fitz.Rect(30, 25, page_width - 30, 70),
                    "CERTIFICACIÓN DE INTEGRIDAD Y DEPÓSITO DE EVIDENCIA IPFS\nSTAR-DOC LEGALTECH PLATFORM",
                    fontsize=11,
                    fontname="Helvetica-Bold",
                    color=(1, 1, 1),
                    align=fitz.TEXT_ALIGN_CENTER
                )
                
                # Caja de metadatos de auditoría
                meta_rect = fitz.Rect(20, 95, page_width - 20, 230)
                cert_page.draw_rect(meta_rect, color=(0.85, 0.85, 0.85), width=1)
                
                from datetime import timedelta
                colombia_now = datetime.utcnow() - timedelta(hours=5)
                meta_text = (
                    f"• Nombre del Archivo: {os.path.basename(pdf_path)}\n"
                    f"• Hash SHA-256 Original: {sha256_original}\n"
                    f"• Dirección IPFS (CID): {cid_original}\n"
                    f"• Gateway de Acceso Seguro: https://ipfs.io/ipfs/{cid_original}\n"
                    f"• Fecha de Certificación: {colombia_now.strftime('%Y-%m-%d %H:%M:%S COT')}"
                )
                cert_page.insert_textbox(
                    fitz.Rect(30, 105, page_width - 30, 220),
                    meta_text,
                    fontsize=9.0,
                    fontname="Helvetica",
                    color=(0.1, 0.1, 0.1)
                )
                
                # Dibujar QR centrado en la página
                qr_rect = fitz.Rect((page_width - 150) / 2, 260, (page_width + 150) / 2, 410)
                cert_page.insert_image(qr_rect, stream=img_byte_arr.read())
                
                cert_page.insert_textbox(
                    fitz.Rect(30, 420, page_width - 30, 440),
                    "Escanee el código QR para verificar la integridad del documento original en la red IPFS",
                    fontsize=8,
                    fontname="Helvetica-Bold",
                    color=(0.3, 0.3, 0.3),
                    align=fitz.TEXT_ALIGN_CENTER
                )
                
                # Declaración de validez legal en la parte inferior
                cert_page.draw_rect(fitz.Rect(20, page_height - 140, page_width - 20, page_height - 40), color=None, fill=(0.95, 0.95, 0.98), overlay=True)
                cert_page.draw_rect(fitz.Rect(20, page_height - 140, page_width - 20, page_height - 40), color=(0.85, 0.85, 0.85), width=1)
                
                disclaimer = (
                    "CERTIFICADO DE DEPOSITARIO DE EVIDENCIA DIGITAL:\n"
                    "Este documento ha sido procesado por la plataforma STAR-DOC y anclado en la red de almacenamiento descentralizada IPFS. "
                    "El hash criptográfico SHA-256 garantiza que el contenido del documento es inalterable, constituyendo prueba en los términos de la "
                    "Ley 527 de 1999 sobre Comercio Electrónico y Firmas Digitales en la República de Colombia. "
                    "Cualquier modificación posterior al documento invalidará el hash original y la firma digital asociada."
                )
                cert_page.insert_textbox(
                    fitz.Rect(30, page_height - 135, page_width - 30, page_height - 45),
                    disclaimer,
                    fontsize=7.5,
                    fontname="Helvetica",
                    color=(0.3, 0.3, 0.3)
                )
                
            temp_path = f"{pdf_path}.tmp"
            doc.save(temp_path)
            doc.close()
            os.replace(temp_path, pdf_path)
            
        import asyncio
        from datetime import datetime
        await asyncio.to_thread(_sync_stamp)

    @staticmethod
    async def _stamp_docx_with_text(docx_path: str, cid_original: str, sha256_original: str):
        """
        Añade una sección de certificación de evidencia al final del archivo DOCX.
        """
        def _sync_stamp():
            from docx import Document
            from docx.shared import Pt, RGBColor
            
            doc = Document(docx_path)
            
            # Párrafo divisor
            p_div = doc.add_paragraph()
            p_div.paragraph_format.space_before = Pt(24)
            p_div.paragraph_format.space_after = Pt(8)
            run_div = p_div.add_run("__________________________________________________________________")
            run_div.font.color.rgb = RGBColor(148, 163, 184) # slate-400
            run_div.font.size = Pt(9)
            
            # Título
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_after = Pt(6)
            run_title = p_title.add_run("CERTIFICACION DE EVIDENCIA DIGITAL (STAR-DOC)")
            run_title.bold = True
            run_title.font.name = "Arial"
            run_title.font.size = Pt(10)
            run_title.font.color.rgb = RGBColor(15, 23, 42) # Slate-900
            
            # Metadatos del sello
            p_meta = doc.add_paragraph()
            p_meta.paragraph_format.line_spacing = 1.15
            p_meta.paragraph_format.space_after = Pt(4)
            
            # Hash SHA256
            run_sha = p_meta.add_run("Hash SHA-256 Original (Limpio): ")
            run_sha.bold = True
            run_sha.font.size = Pt(8.5)
            run_sha.font.color.rgb = RGBColor(71, 85, 105) # Slate-600
            run_sha_val = p_meta.add_run(f"{sha256_original}\n")
            run_sha_val.font.size = Pt(8.5)
            run_sha_val.font.name = "Courier New"
            
            # Dirección IPFS
            run_cid = p_meta.add_run("Direccion IPFS (Original): ")
            run_cid.bold = True
            run_cid.font.size = Pt(8.5)
            run_cid.font.color.rgb = RGBColor(71, 85, 105) # Slate-600
            run_cid_val = p_meta.add_run(f"https://ipfs.io/ipfs/{cid_original}\n")
            run_cid_val.font.size = Pt(8.5)
            run_cid_val.font.name = "Courier New"
            
            # Texto explicativo
            run_desc = p_meta.add_run(
                "Este documento y sus firmas asociadas han sido anclados en un registro inmutable descentralizado. "
                "Cualquier alteracion fisica o logica del contenido invalidara el Hash SHA-256 original arriba indicado."
            )
            run_desc.italic = True
            run_desc.font.size = Pt(8)
            run_desc.font.color.rgb = RGBColor(100, 116, 139) # Slate-500
            
            doc.save(docx_path)
            
        import asyncio
        await asyncio.to_thread(_sync_stamp)

    @staticmethod
    async def pack_audit(
        audit_name: str,
        document_ids: list[int],
        session: AsyncSession
    ) -> "IPFSAudit":
        """
        Agrupa múltiples documentos registrados en un único directorio IPFS (Merkle DAG).
        """
        from sqlmodel import select
        from app.models.ipfs_audit import IPFSAudit
        import re
        from datetime import datetime

        # 1. Recuperar los registros de documentos de la base de datos
        stmt = select(DocumentIPFS).where(DocumentIPFS.id.in_(document_ids))
        res = await session.execute(stmt)
        docs = res.scalars().all()
        
        if not docs:
            raise ValueError("No se encontraron documentos válidos para empaquetar.")

        # 2. Generar nombre de carpeta único
        sanitized_name = re.sub(r'[^a-zA-Z0-9_\-]', '_', audit_name)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        mfs_dir_name = f"{sanitized_name}_{timestamp}"
        mfs_dir_path = f"/audits/{mfs_dir_name}"

        # 3. Crear el directorio en Kubo MFS
        await IPFSService.create_mfs_directory(mfs_dir_path)

        # 4. Copiar cada documento al directorio MFS
        for doc in docs:
            # Para evitar colisiones de nombres si se llaman igual, usamos prefix de ID
            filename_in_dir = f"{doc.id}_{doc.original_filename}"
            # Asegurarse de que no tenga caracteres prohibidos en rutas MFS
            filename_in_dir = re.sub(r'[\\/*?:"<>|]', '_', filename_in_dir)
            dest_path = f"{mfs_dir_path}/{filename_in_dir}"
            
            try:
                await IPFSService.copy_to_mfs(f"/ipfs/{doc.ipfs_cid}", dest_path)
            except Exception as e:
                logger.error(f"Error copiando {doc.ipfs_cid} a {dest_path}: {e}")
                raise RuntimeError(f"Fallo al agregar documento {doc.original_filename} al paquete: {e}")

        # 5. Obtener el CID del directorio
        stat_result = await IPFSService.stat_mfs_path(mfs_dir_path)
        dir_cid = stat_result["Hash"]

        # 6. Pinear el CID de manera persistente en Kubo
        await IPFSService.pin_in_kubo(dir_cid)

        # Clasificación resultante es la más restrictiva
        # chain_of_custody > confidential > public
        classifications = [doc.classification for doc in docs]
        if "chain_of_custody" in classifications:
            final_class = "chain_of_custody"
        elif "confidential" in classifications:
            final_class = "confidential"
        else:
            final_class = "public"

        # 7. Redundancia en Pinata si JWT está configurado (Estrategia híbrida para proteger la cuota de 500 archivos)
        if settings.PINATA_JWT and final_class != "public":
            try:
                await IPFSService.pin_by_cid_pinata(dir_cid, mfs_dir_name)
                logger.info(f"Paquete de auditoría {mfs_dir_name} pineado en Pinata: {dir_cid}")
            except Exception as e:
                logger.warning(f"No se pudo pinear el paquete en Pinata: {e}")
        else:
            logger.info("Evitando pin en Pinata para auditoria publica (Estrategia Hibrida para proteger limite gratuito)")

        # 8. Registrar la auditoría en la base de datos
        audit_record = IPFSAudit(
            name=audit_name,
            ipfs_cid=dir_cid,
            document_ids=document_ids,
            classification=final_class
        )
        session.add(audit_record)
        await session.commit()
        await session.refresh(audit_record)

        logger.info(f"Auditoría empaquetada y registrada: ID={audit_record.id}, CID={dir_cid}")
        return audit_record

    @staticmethod
    async def get_document_by_cid(cid_or_id: str, session: AsyncSession) -> Optional[DocumentIPFS]:
        """Busca un documento en la base de datos por CID (estampado o original) o ID."""
        from sqlmodel import select
        # Intentar buscar por ID si es un entero
        try:
            doc_id = int(cid_or_id)
            stmt = select(DocumentIPFS).where(DocumentIPFS.id == doc_id)
            res = await session.execute(stmt)
            doc = res.scalar_one_or_none()
            if doc:
                return doc
        except ValueError:
            pass
        
        # Buscar por CID
        stmt = select(DocumentIPFS).where(
            (DocumentIPFS.ipfs_cid == cid_or_id) | (DocumentIPFS.ipfs_cid_original == cid_or_id)
        )
        res = await session.execute(stmt)
        return res.scalar_one_or_none()

    @staticmethod
    async def list_documents_records(limit: int, offset: int, session: AsyncSession) -> list[DocumentIPFS]:
        """Lista los documentos registrados en base de datos ordenados descendentemente."""
        from sqlmodel import select
        stmt = select(DocumentIPFS).order_by(DocumentIPFS.created_at.desc()).offset(offset).limit(limit)
        res = await session.execute(stmt)
        return res.scalars().all()

    @staticmethod
    async def get_audit_by_cid_or_id(id_or_cid: str, session: AsyncSession) -> Optional[IPFSAudit]:
        """Busca un expediente de auditoría por ID o CID."""
        from sqlmodel import select
        try:
            audit_id = int(id_or_cid)
            stmt = select(IPFSAudit).where(IPFSAudit.id == audit_id)
            res = await session.execute(stmt)
            audit = res.scalar_one_or_none()
            if audit:
                return audit
        except ValueError:
            pass

        stmt = select(IPFSAudit).where(IPFSAudit.ipfs_cid == id_or_cid)
        res = await session.execute(stmt)
        return res.scalar_one_or_none()

    @staticmethod
    async def list_audits_records(limit: int, offset: int, session: AsyncSession) -> list[IPFSAudit]:
        """Lista los paquetes de auditoría registrados en la base de datos."""
        from sqlmodel import select
        stmt = select(IPFSAudit).order_by(IPFSAudit.created_at.desc()).offset(offset).limit(limit)
        res = await session.execute(stmt)
        return res.scalars().all()

    @staticmethod
    async def get_access_logs(cids: list[str], session: AsyncSession):
        """Busca logs de acceso para una lista de CIDs."""
        from sqlmodel import select
        from app.models.document_access_log import DocumentAccessLog
        stmt = select(DocumentAccessLog).where(DocumentAccessLog.ipfs_cid.in_(cids)).order_by(DocumentAccessLog.accessed_at.desc())
        res = await session.execute(stmt)
        return res.scalars().all()

    @staticmethod
    async def create_webhook_subscription(name: str, url: str, secret: str, events: list[str], session: AsyncSession):
        """Crea y registra una nueva suscripción de webhook en base de datos."""
        from app.models.webhook_subscription import WebhookSubscription
        sub = WebhookSubscription(
            name=name,
            url=url,
            secret=secret,
            events=events,
            is_active=True
        )
        session.add(sub)
        await session.commit()
        await session.refresh(sub)
        return sub

    @staticmethod
    async def list_webhooks(session: AsyncSession):
        """Lista todas las suscripciones a Webhooks activas."""
        from sqlmodel import select
        from app.models.webhook_subscription import WebhookSubscription
        stmt = select(WebhookSubscription)
        res = await session.execute(stmt)
        return res.scalars().all()

    @staticmethod
    async def delete_webhook_subscription(webhook_id: int, session: AsyncSession) -> bool:
        """Elimina una suscripción de webhook por su ID."""
        from sqlmodel import select
        from app.models.webhook_subscription import WebhookSubscription
        stmt = select(WebhookSubscription).where(WebhookSubscription.id == webhook_id)
        res = await session.execute(stmt)
        sub = res.scalar_one_or_none()
        if not sub:
            return False
        await session.delete(sub)
        await session.commit()
        return True

    @staticmethod
    async def get_ipns_key(key_name: str, session: AsyncSession):
        """Obtiene la clave IPNS registrada en la base de datos."""
        from sqlmodel import select
        from app.models.ipns_key import IPNSKey
        stmt = select(IPNSKey).where(IPNSKey.key_name == key_name)
        res = await session.execute(stmt)
        return res.scalar_one_or_none()

    @staticmethod
    async def save_ipns_key(key_name: str, ipns_id: str, cid: Optional[str], session: AsyncSession):
        """Guarda o actualiza el CID de una clave IPNS."""
        from app.models.ipns_key import IPNSKey
        from datetime import datetime, timezone
        db_key = await IPFSIntegrationService.get_ipns_key(key_name, session)
        if db_key:
            if cid:
                db_key.current_cid = cid
                db_key.last_published_at = datetime.now(timezone.utc).replace(tzinfo=None)
            session.add(db_key)
            await session.commit()
            await session.refresh(db_key)
        else:
            db_key = IPNSKey(
                key_name=key_name,
                ipns_id=ipns_id,
                current_cid=cid,
                last_published_at=datetime.now(timezone.utc).replace(tzinfo=None) if cid else None,
                is_active=True
            )
            session.add(db_key)
            await session.commit()
            await session.refresh(db_key)
        return db_key

    @staticmethod
    async def list_ipns_keys(session: AsyncSession):
        """Lista todas las claves IPNS registradas en la base de datos."""
        from sqlmodel import select
        from app.models.ipns_key import IPNSKey
        stmt = select(IPNSKey)
        res = await session.execute(stmt)
        return res.scalars().all()

    @staticmethod
    async def unpin_document_record(cid: str, session: AsyncSession) -> bool:
        """Actualiza el estado de pinning local en base de datos al despinear."""
        doc = await IPFSIntegrationService.get_document_by_cid(cid, session)
        if not doc:
            return False
        doc.pinned_kubo = False
        session.add(doc)
        await session.commit()
        return True

    @staticmethod
    async def sync_pinata_record(cid: str, session: AsyncSession) -> bool:
        """Actualiza el estado de pinning en Pinata en base de datos al sincronizar."""
        doc = await IPFSIntegrationService.get_document_by_cid(cid, session)
        if not doc:
            return False
        doc.pinned_pinata = True
        session.add(doc)
        await session.commit()
        return True

    @staticmethod
    async def register_access_log(
        cid: str,
        action: str,
        ip_address: Optional[str],
        user_agent: Optional[str],
        user_id: Optional[int],
        username: str,
        session: AsyncSession
    ):
        """Registra una entrada de acceso a documentos en PostgreSQL para la cadena de custodia inmutable."""
        from app.models.document_access_log import DocumentAccessLog
        log_entry = DocumentAccessLog(
            ipfs_cid=cid,
            user_id=user_id,
            username=username,
            action=action,
            ip_address=ip_address,
            user_agent=user_agent
        )
        session.add(log_entry)
        await session.commit()
