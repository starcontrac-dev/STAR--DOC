"""
Servicio para la comparación inteligente de documentos (Diff IA) en STAR-DOC.
Extrae texto de múltiples formatos, genera diffs visuales en HTML y
realiza análisis de riesgos jurídicos usando Gemini API.
"""
import os
import logging
import difflib
import json
from typing import Dict, Any, Optional

import docx
import fitz  # PyMuPDF
from app.services.ai_service import ai_service

logger = logging.getLogger(__name__)

class DiffService:
    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extrae el texto de un archivo según su extensión (.md, .docx, .pdf).
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
            
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.md' or ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
                    
            elif ext == '.docx':
                doc = docx.Document(file_path)
                text_parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text)
                return "\n".join(text_parts)
                
            elif ext == '.pdf':
                doc = fitz.open(file_path)
                text_parts = []
                for page in doc:
                    text_parts.append(page.get_text())
                return "\n".join(text_parts)
                
            else:
                # Fallback para archivos sin extensión clara, intentar como texto
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
        except Exception as e:
            logger.error(f"Error extrayendo texto de {file_path}: {e}")
            raise ValueError(f"No se pudo extraer texto del archivo: {str(e)}")

    @staticmethod
    def generate_html_diff(text_a: str, text_b: str) -> str:
        """
        Calcula la diferencia línea por línea y genera un bloque de código HTML
        estilizado con Tailwind CSS para visualización premium (estilo GitHub).
        """
        lines_a = text_a.splitlines()
        lines_b = text_b.splitlines()
        
        # ndiff proporciona diff detallado con marcas '+', '-' y ' '
        diff = list(difflib.ndiff(lines_a, lines_b))
        
        html = [
            '<div class="diff-container font-mono text-xs overflow-y-auto p-4 bg-[#0a0f1d] text-gray-300 rounded-xl border border-white/10 max-h-[550px] custom-scrollbar space-y-1">'
        ]
        
        for line in diff:
            if not line:
                continue
            prefix = line[:2]
            content = line[2:]
            
            # Escapar caracteres HTML peligrosos
            content_safe = content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            if prefix == '- ':
                html.append(
                    f'<div class="bg-red-950/30 text-red-300 border-l-2 border-red-500/80 px-3 py-1 rounded-sm whitespace-pre-wrap"><span class="select-none font-bold text-red-500 mr-2">-</span>{content_safe}</div>'
                )
            elif prefix == '+ ':
                html.append(
                    f'<div class="bg-green-950/30 text-green-300 border-l-2 border-green-500/80 px-3 py-1 rounded-sm whitespace-pre-wrap"><span class="select-none font-bold text-green-500 mr-2">+</span>{content_safe}</div>'
                )
            elif prefix == '  ':
                html.append(
                    f'<div class="px-3 py-0.5 whitespace-pre-wrap text-gray-400/90"><span class="select-none text-gray-700 mr-2">&nbsp;</span>{content_safe}</div>'
                )
            # Ignoramos la línea de guía de cambios '? ' de ndiff por claridad visual en el visor dual
            
        html.append('</div>')
        return "\n".join(html)

    @staticmethod
    async def analyze_legal_risk_with_gemini(text_a: str, text_b: str) -> Dict[str, Any]:
        """
        Envía los textos de las dos versiones a Gemini y solicita un análisis legal
        de diferencias estructurado en JSON (enfocado en legislación de Colombia).
        """
        prompt = f"""
Actúa como un Auditor Legal Senior y Abogado Colombiano experto. Tu labor consiste en realizar una auditoría de diferencias contractuales.
Compara las siguientes dos versiones de un documento legal y evalúa si las modificaciones introducen riesgos, alteran el equilibrio contractual o tienen implicaciones legales bajo el derecho colombiano.

VERSION ORIGINAL (A):
---
{text_a[:12000]}
---

VERSION MODIFICADA (B):
---
{text_b[:12000]}
---

Tu respuesta debe ser exclusivamente un JSON válido en español que analice el impacto legal y el nivel de riesgo de los cambios. No agregues texto introductorio o conclusivo fuera del JSON. El JSON debe cumplir con esta estructura:
{{
    "resumen_cambios": "Resumen global explicativo de las modificaciones principales.",
    "evaluacion_riesgo": "Explicación integral del riesgo legal global detectado bajo la normativa colombiana.",
    "nivel_riesgo": "Bajo" | "Medio" | "Alto",
    "modificaciones_detectadas": [
        {{
            "seccion": "Cláusula o sección afectada",
            "descripcion": "Detalle breve del cambio textual",
            "implicacion_legal": "Explicación jurídica de cómo altera este cambio la posición del cliente según el derecho comercial/civil colombiano (ej. responsabilidad, indemnidad, mora, prórrogas).",
            "riesgo": "Bajo" | "Medio" | "Alto"
        }}
    ],
    "recomendaciones": [
        "Recomendación o acción sugerida para blindar el documento contra este cambio...",
        "Otra recomendación..."
    ]
}}
"""
        payload = {
            "contents": [
                {"role": "user", "parts": [{"text": prompt}]}
            ],
            "generationConfig": {
                "responseMimeType": "application/json",
                "temperature": 0.2
            }
        }
        
        try:
            response = await ai_service.generate_content(payload=payload, model="gemini-2.5-flash")
            candidate = response.get("candidates", [{}])[0]
            text_response = candidate.get("content", {}).get("parts", [{}])[0].get("text", "")
            
            # Intentar parsear el JSON retornado
            return json.loads(text_response)
        except Exception as e:
            logger.error(f"Error llamando a Gemini para Diff IA: {e}")
            return {
                "resumen_cambios": "No se pudo generar el análisis automático de la IA debido a un error técnico.",
                "evaluacion_riesgo": "Desconocido. Error al invocar la API del modelo de lenguaje.",
                "nivel_riesgo": "Medio",
                "modificaciones_detectadas": [],
                "recomendaciones": ["Por favor verifique las diferencias visuales en el panel manualmente."]
            }
